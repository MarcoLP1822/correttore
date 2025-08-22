# core/document_handler.py
"""
Gestione dei documenti Word con validazione, backup e manipolazione sicura.
Separazione completa delle responsabilità I/O dai processi di correzione.
"""

import logging
import shutil
import zipfile
from pathlib import Path
from typing import List, Iterable, Optional, Tuple
from dataclasses import dataclass

from docx import Document
from docx.text.paragraph import Paragraph

from src.core.validation import DocumentValidator, ValidationResult, create_backup
from config.settings import get_validation_config

logger = logging.getLogger(__name__)


class DocumentLoadError(Exception):
    """Errore durante il caricamento del documento"""
    pass


class DocumentSaveError(Exception):
    """Errore durante il salvataggio del documento"""
    pass


class DocumentValidationError(Exception):
    """Errore durante la validazione del documento"""
    pass

@dataclass
class DocumentInfo:
    """Informazioni su un documento"""
    path: Path
    total_paragraphs: int
    total_characters: int
    needs_correction_count: int
    validation_result: ValidationResult
    backup_path: Optional[Path] = None

class DocumentHandler:
    """Gestore centralizzato per operazioni sui documenti Word"""
    
    def __init__(self):
        self.validator = DocumentValidator()
        self.config = get_validation_config()
        
    def load_document(self, doc_path: Path, create_backup_copy: bool = True) -> Tuple[Document, DocumentInfo]:  # type: ignore
        """
        Carica un documento Word con validazione e backup opzionale.
        
        Returns:
            Tuple[Document, DocumentInfo]: Documento caricato e sue informazioni
        """
        logger.info(f"📁 Loading document: {doc_path.name}")
        
        # 1. Validazione preliminare
        validation_result = self.validator.validate_before_processing(doc_path)
        if not validation_result.is_valid:
            critical_issues = [issue for issue in validation_result.issues 
                             if any(word in issue.lower() for word in ['error', 'cannot', 'failed', 'invalid'])]
            if critical_issues:
                raise ValueError(f"Document validation failed: {critical_issues}")
            else:
                logger.warning("⚠️  Document has minor issues but proceeding")
                for issue in validation_result.issues:
                    logger.warning(f"   • {issue}")
        
        # 2. Backup se richiesto
        backup_path = None
        if create_backup_copy and self.config.backup_enabled:
            backup_path = create_backup(doc_path)
            logger.info(f"💾 Backup created: {backup_path.name}")
        
        # 3. Caricamento documento
        try:
            doc = Document(str(doc_path))
        except Exception as e:
            logger.error(f"❌ Failed to load document: {e}")
            raise RuntimeError(f"Cannot load document: {e}")
        
        # 4. Analisi documento
        all_paragraphs = list(self._iter_all_paragraphs(doc))
        needs_correction = [p for p in all_paragraphs if self._paragraph_needs_correction(p)]
        total_chars = sum(len(p.text) for p in all_paragraphs)
        
        doc_info = DocumentInfo(
            path=doc_path,
            total_paragraphs=len(all_paragraphs),
            total_characters=total_chars,
            needs_correction_count=len(needs_correction),
            validation_result=validation_result,
            backup_path=backup_path
        )
        
        logger.info(f"✅ Document loaded: {doc_info.total_paragraphs} paragraphs, {doc_info.needs_correction_count} need correction")
        return doc, doc_info
    
    def save_document(self, doc: Document, output_path: Path, validate_after_save: bool = True) -> bool:  # type: ignore
        """
        Salva un documento con validazione opzionale post-save.
        
        Returns:
            bool: True se salvato con successo
        """
        logger.info(f"💾 Saving document: {output_path.name}")
        
        try:
            # Assicurati che la directory esista
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Salva il documento
            doc.save(output_path)
            
            # Validazione post-save se richiesta
            if validate_after_save:
                final_validation = self.validator.validate_before_processing(output_path)
                if not final_validation.is_valid:
                    logger.error("❌ Saved document validation failed!")
                    return False
                else:
                    logger.info("✅ Saved document validation passed")
            
            logger.info(f"✅ Document saved successfully: {output_path.name}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Failed to save document: {e}")
            return False
    
    def restore_from_backup(self, backup_path: Path, target_path: Path) -> bool:
        """Ripristina un documento da backup"""
        return self.validator.restore_from_backup(backup_path, target_path)
    
    def extract_all_paragraphs(self, doc: Document) -> List[Paragraph]:  # type: ignore
        """Estrae tutti i paragrafi dal documento (corpo, header, footer, note, caselle)"""
        return list(self._iter_all_paragraphs(doc))
    
    def extract_footnotes_xml(self, docx_path: Path) -> Optional[str]:
        """
        Estrae il contenuto XML delle note a piè di pagina.
        Ritorna None se non ci sono note.
        """
        try:
            with zipfile.ZipFile(docx_path, "r") as zf:
                footnotes_file = "word/footnotes.xml"
                if footnotes_file in zf.namelist():
                    return zf.read(footnotes_file).decode('utf-8')
            return None
        except Exception as e:
            logger.error(f"❌ Failed to extract footnotes XML: {e}")
            return None
    
    def update_footnotes_xml(self, docx_path: Path, footnotes_xml: str) -> bool:
        """
        Aggiorna il file footnotes.xml nel documento.
        Ritorna True se l'operazione è riuscita.
        """
        try:
            # Estrai tutto in una directory temporanea
            temp_dir = docx_path.parent / "_temp_docx_update"
            temp_dir.mkdir(exist_ok=True)
            
            with zipfile.ZipFile(docx_path, "r") as zf:
                zf.extractall(temp_dir)
            
            # Aggiorna footnotes.xml
            footnotes_file = temp_dir / "word" / "footnotes.xml"
            footnotes_file.write_text(footnotes_xml, encoding='utf-8')
            
            # Ricrea il .docx
            temp_docx = docx_path.with_suffix('.tmp')
            with zipfile.ZipFile(temp_docx, 'w') as zf:
                for file_path in temp_dir.rglob('*'):
                    if file_path.is_file():
                        arcname = file_path.relative_to(temp_dir)
                        zf.write(file_path, arcname)
            
            # Sostituisci l'originale
            shutil.move(temp_docx, docx_path)
            shutil.rmtree(temp_dir)
            
            logger.info("✅ Footnotes XML updated successfully")
            return True
            
        except Exception as e:
            logger.error(f"❌ Failed to update footnotes XML: {e}")
            # Cleanup in caso di errore
            if temp_dir.exists():
                shutil.rmtree(temp_dir, ignore_errors=True)
            if temp_docx.exists():
                temp_docx.unlink(missing_ok=True)
            return False
    
    def cleanup_old_backups(self, backup_dir: Path, retention_days: Optional[int] = None) -> int:
        """
        Pulisce backup vecchi oltre il periodo di retention.
        
        Returns:
            int: Numero di backup eliminati
        """
        if retention_days is None:
            retention_days = self.config.backup_retention_days
            
        if not backup_dir.exists():
            return 0
            
        from datetime import datetime, timedelta
        cutoff_date = datetime.now() - timedelta(days=retention_days)
        
        removed_count = 0
        for backup_file in backup_dir.glob("*_backup.*"):
            try:
                if datetime.fromtimestamp(backup_file.stat().st_mtime) < cutoff_date:
                    backup_file.unlink()
                    # Rimuovi anche metadati se esistono
                    metadata_file = backup_file.with_suffix('.backup.json')
                    if metadata_file.exists():
                        metadata_file.unlink()
                    removed_count += 1
            except Exception as e:
                logger.warning(f"⚠️  Could not remove old backup {backup_file.name}: {e}")
        
        if removed_count > 0:
            logger.info(f"🧹 Cleaned up {removed_count} old backups")
        
        return removed_count
    
    # Metodi privati
    
    def _iter_all_paragraphs(self, doc: Document) -> Iterable[Paragraph]:  # type: ignore
        """Itera su tutti i paragrafi del documento"""
        # Corpo principale
        yield from self._iter_body_paragraphs(doc)
        # Note a piè di pagina
        yield from self._iter_footnote_paragraphs(doc)
        # Header e footer
        yield from self._iter_header_footer_paragraphs(doc)
        # Caselle di testo
        yield from self._iter_textbox_paragraphs(doc)
    
    def _iter_body_paragraphs(self, container) -> Iterable[Paragraph]:
        """Itera sui paragrafi del corpo (incluse tabelle)"""
        for para in container.paragraphs:
            yield para
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    yield from self._iter_body_paragraphs(cell)
    
    def _iter_footnote_paragraphs(self, doc: Document) -> Iterable[Paragraph]:  # type: ignore
        """Itera sui paragrafi delle note a piè di pagina"""
        footnotes_part = getattr(doc.part, "footnotes_part", None)
        if footnotes_part:
            for footnote in footnotes_part.footnotes:
                for para in footnote.paragraphs:
                    yield para
    
    def _iter_header_footer_paragraphs(self, doc: Document) -> Iterable[Paragraph]:  # type: ignore
        """Itera sui paragrafi di header e footer"""
        for section in doc.sections:
            for hf in (section.header, section.footer):
                if hf:
                    yield from self._iter_body_paragraphs(hf)
    
    def _iter_textbox_paragraphs(self, doc: Document) -> Iterable[Paragraph]:  # type: ignore
        """Itera sui paragrafi delle caselle di testo"""
        parts = [doc.part]
        for section in doc.sections:
            parts.extend([section.header.part, section.footer.part])
        footnotes_part = getattr(doc.part, "footnotes_part", None)
        if footnotes_part:
            parts.append(footnotes_part)
            
        for part in parts:
            root = part._element
            for txbx in root.xpath('.//*[local-name()="txbxContent"]'):
                for p_el in txbx.xpath('.//*[local-name()="p"]'):
                    yield Paragraph(p_el, part)
    
    def _paragraph_needs_correction(self, paragraph: Paragraph) -> bool:
        """Determina se un paragrafo ha bisogno di correzione"""
        if not paragraph.text.strip():
            return False
            
        # Evita paragrafi con equazioni matematiche
        if self._has_math(paragraph):
            return False
            
        # Applica euristica per determinare se ha errori
        # (questo potrebbe essere estratto in un servizio separato)
        return self._has_potential_errors(paragraph.text)
    
    def _has_math(self, paragraph: Paragraph) -> bool:
        """Controlla se il paragrafo contiene equazioni matematiche"""
        return bool(paragraph._p.xpath(".//*[local-name()='oMath' or local-name()='oMathPara']"))
    
    def _has_potential_errors(self, text: str) -> bool:
        """Euristica veloce per determinare se un testo potrebbe avere errori"""
        import re
        
        # Pattern di errori comuni
        error_patterns = [
            r"( {2,})",                     # spazi doppi
            r"([,;:.!?]\S)",               # punteggiatura+lettera
            r"(«\s)|(\s»)",                # spazio dopo « o prima »
            r"CAPPIT",                     # "CAPPITOLO" mal scritto
            r"\bpò\b",                     # "pò" invece di "po'"
        ]
        
        for pattern in error_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
                
        return False


# Funzione di convenienza per uso rapido
def load_document_safe(doc_path: Path) -> Tuple[Document, DocumentInfo]:  # type: ignore
    """Carica un documento in modo sicuro con validazione e backup"""
    handler = DocumentHandler()
    return handler.load_document(doc_path)
