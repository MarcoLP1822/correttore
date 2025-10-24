#!/usr/bin/env python3
"""
Crea un documento di test per verificare il correttore
"""

from docx import Document

def create_test_document():
    doc = Document()
    
    # Aggiungi paragrafi con errori intenzionali
    doc.add_paragraph('Ciao, questo è un testo di prova per il correttore automatico. Il gatto magna il pesce.')
    doc.add_paragraph('Questo paragrafo contiene degli errore di battitura intenzionnali per testare il sistema.')
    doc.add_paragraph('La machina è molto veloce e il cane corri nel parco.')
    
    # Salva il documento
    doc.save('test_input.docx')
    print('✅ File di test creato: test_input.docx')

if __name__ == "__main__":
    create_test_document()
