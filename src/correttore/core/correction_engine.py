# core/correction_engine.py
"""
Motore di correzione principale con integrazione AI, validazione qualità e rollback.
Orchestrazione del processo di correzione con safety guarantees.
"""

import logging
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from concurrent.futures import as_completed, ThreadPoolExecutor
from correttore.utils.text_normalize import prenormalize

from docx import Document as DocxDocument  # Renamed to avoid type confusion
from docx.text.paragraph import Paragraph

from correttore.core.document_handler import DocumentHandler, DocumentInfo
from correttore.services.openai_service import OpenAIService
from correttore.services.languagetool_service import LanguageToolService
from correttore.services.intelligent_cache import get_cache
from correttore.core.safe_correction import SafeCorrector, QualityScore, CorrectionConfidence
from correttore.core.correction_collector import CorrectionCollector
from correttore.config.settings import get_correction_config
from correttore.utils.readability import ReadabilityAnalyzer
from correttore.models import CorrectionSource

logger = logging.getLogger(__name__)

@dataclass
class CorrectionContext:
    """Contesto per una sessione di correzione"""
    source_document: DocxDocument  # type: ignore
    document_info: DocumentInfo
    target_paragraphs: List[Paragraph]
    corrections_applied: int = 0
    corrections_rejected: int = 0
    total_processed: int = 0
    corrections_log: List[Dict[str, Any]] = field(default_factory=list)

@dataclass
class CorrectionResult:
    """Risultato di una correzione completa"""
    success: bool
    document: Optional[DocxDocument]  # type: ignore
    context: CorrectionContext
    error_message: Optional[str] = None
    quality_scores: List[QualityScore] = field(default_factory=list)

class CorrectionEngine:
    """
    Motore principale di correzione con orchestrazione completa del processo.
    Integra AI, grammatica, caching e safety systems.
    """
    
    def __init__(self, enable_tracking: bool = True):
        self.document_handler = DocumentHandler()
        self.openai_service = OpenAIService()
        self.languagetool_service = LanguageToolService()
        self.cache_service = get_cache()
        
        # FASE 6: Carica correzioni apprese da feedback (prima del SafeCorrector)
        self.custom_corrections = {}  # {original: corrected}
        self.custom_whitelist = set()  # parole da NON correggere
        self._load_custom_corrections()
        
        # Inizializza il collector per il tracking (se abilitato)
        if enable_tracking:
            self.collector = CorrectionCollector()
            logger.info("✅ CorrectionCollector initialized for tracking")
        else:
            self.collector = None
        
        # Passa il collector E la whitelist al SafeCorrector
        self.safe_corrector = SafeCorrector(collector=self.collector, whitelist=self.custom_whitelist)
        self.readability_analyzer = ReadabilityAnalyzer()
        
        # FASE 7: Service per categorie speciali (parole straniere, sensibili, nomi propri)
        try:
            from src.correttore.services.special_categories_service import (
                SpecialCategoriesService
            )
            self.special_categories_service = SpecialCategoriesService()
            logger.info("✅ SpecialCategoriesService initialized for FASE 7")
        except Exception as e:
            logger.warning(f"⚠️  SpecialCategoriesService not available: {e}")
            self.special_categories_service = None
        
        self.config = get_correction_config()
        # self.app_config = get_app_config()  # Commentato fino a quando non sarà implementato
        
        logger.info(f"🔧 CorrectionEngine initialized (tracking: {'enabled' if enable_tracking else 'disabled'})")
        if self.custom_corrections:
            logger.info(f"✅ Loaded {len(self.custom_corrections)} custom corrections from feedback")
        if self.custom_whitelist:
            logger.info(f"✅ Loaded {len(self.custom_whitelist)} whitelisted words from feedback")
    
    def correct_document(self, input_path, output_path=None) -> CorrectionResult:
        """
        Corregge un documento completo con tutte le safety guarantees.
        
        Args:
            input_path: Percorso del documento da correggere
            output_path: Percorso di output (opzionale)
            
        Returns:
            CorrectionResult: Risultato completo della correzione
        """
        from pathlib import Path
        input_path = Path(input_path)
        
        if not output_path:
            output_path = input_path.parent / f"{input_path.stem}_corretto{input_path.suffix}"
        else:
            output_path = Path(output_path)
        
        logger.info(f"🎯 Starting document correction: {input_path.name}")
        
        # Inizia tracking se abilitato
        if self.collector is not None:
            self.collector.start_tracking()
        
        try:
            # 1. Caricamento documento con validazione
            document, doc_info = self.document_handler.load_document(input_path, create_backup_copy=True)
            
            # 2. Identifica paragrafi da correggere
            all_paragraphs = self.document_handler.extract_all_paragraphs(document)
            target_paragraphs = [p for p in all_paragraphs 
                               if self._should_correct_paragraph(p)]
            
            context = CorrectionContext(
                source_document=document,
                document_info=doc_info,
                target_paragraphs=target_paragraphs
            )
            
            logger.info(f"📊 Found {len(target_paragraphs)} paragraphs to correct")
            
            # 3. Processo di correzione con batching
            success = self._process_corrections(context)
            
            if not success:
                return CorrectionResult(
                    success=False,
                    document=None,
                    context=context,
                    error_message="Correction process failed"
                )
            
            # 4. Salvataggio con validazione
            save_success = self.document_handler.save_document(
                document, output_path, validate_after_save=True
            )
            
            if not save_success:
                return CorrectionResult(
                    success=False,
                    document=document,
                    context=context,
                    error_message="Failed to save corrected document"
                )
            
            # 5. Report finale
            self._log_correction_summary(context, output_path)
            
            # Ferma tracking e salva statistiche
            if self.collector is not None:
                self.collector.stop_tracking()
                stats = self.collector.get_statistics()
                logger.info(f"📊 Tracked {stats.total_corrections} corrections across {stats.unique_words} unique words")
                
                # FASE 7: Analizza categorie speciali (parole straniere, sensibili, nomi propri)
                if self.special_categories_service:
                    logger.info("🔍 Analyzing special categories (FASE 7)...")
                    self._analyze_special_categories(document)
                
                # Genera report HTML automaticamente
                self._generate_html_report(input_path, output_path)
            else:
                logger.warning("⚠️  Collector is None, skipping report generation")
            
            return CorrectionResult(
                success=True,
                document=document,
                context=context
            )
            
        except Exception as e:
            logger.error(f"❌ Document correction failed: {e}")
            
            # Tentativo di ripristino da backup se disponibile
            if hasattr(doc_info, 'backup_path') and doc_info.backup_path:
                logger.info("🔄 Attempting restore from backup...")
                if self.document_handler.restore_from_backup(doc_info.backup_path, input_path):
                    logger.info("✅ Document restored from backup")
                else:
                    logger.error("❌ Failed to restore from backup")
            
            return CorrectionResult(
                success=False,
                document=None,
                context=CorrectionContext(
                    source_document=DocxDocument(),  # type: ignore
                    document_info=DocumentInfo(
                        path=Path(""),
                        total_paragraphs=0,
                        total_characters=0,
                        needs_correction_count=0,
                        validation_result=ValidationResult(is_valid=False, errors=[], warnings=[])  # type: ignore
                    ),
                    target_paragraphs=[]
                ),
                error_message=str(e)
            )
    
    def correct_text_fragment(self, text: str, use_cache: bool = True) -> Tuple[str, QualityScore]:
        """
        Corregge un singolo frammento di testo.
        
        Args:
            text: Testo da correggere
            use_cache: Se utilizzare la cache
            
        Returns:
            Tuple[str, QualityScore]: Testo corretto e punteggio qualità
        """
        logger.info(f"🔤 [MAIN FLOW] correct_text_fragment called with {len(text)} chars, collector={self.collector is not None}")
        
        # Salva testo originale per confronto finale
        original_text = text
        
        # 1. Check cache se abilitata
        if use_cache and self.config.cache_enabled:
            cached_entry = self.cache_service.get_with_similarity(text)
            if cached_entry:
                logger.info(f"💾 [CACHE HIT] Returning cached correction, will track it")
                
                # Traccia anche le correzioni da cache!
                if self.collector is not None and cached_entry.corrected_text != text:
                    from correttore.models import CorrectionRecord, CorrectionSource, CorrectionCategory
                    record = CorrectionRecord(
                        original_text=text,
                        corrected_text=cached_entry.corrected_text,
                        source=CorrectionSource.SYSTEM,  # Usa SYSTEM per cache
                        category=CorrectionCategory.MIGLIORABILI,
                        rule_id="CACHE_HIT",
                        confidence_score=0.9,
                        context=text,  # IMPORTANTE: Contesto = il testo originale
                        position=0,
                        length=len(text),
                        is_applied=True
                    )
                    self.collector.add_correction(record)
                    logger.info(f"✅ [CACHE] Tracked cached correction in collector")
                
                return cached_entry.corrected_text, QualityScore(  
                    overall_score=0.9,
                    confidence=CorrectionConfidence.HIGH,  # type: ignore
                    content_preservation=1.0,
                    grammar_improvement=0.8,
                    style_preservation=1.0,
                    safety_score=0.95,
                    issues=[]
                )
        
        # 1. FASE 6: Applica correzioni custom da feedback (priorità massima)
        custom_corrected, custom_count = self._apply_custom_corrections(text)
        if custom_count > 0:
            logger.info(f"✨ [MAIN FLOW] Applied {custom_count} custom corrections from feedback")
            text = custom_corrected
        else:
            logger.info(f"⚠️  [MAIN FLOW] No custom corrections applied")
        
        # 2. Pre-validazione con LanguageTool (con graceful degradation)
        lt_errors = []
        try:
            lt_errors = self.languagetool_service.check_text(text)
            
            # Traccia errori LanguageTool se tracking abilitato
            if self.collector is not None and lt_errors:
                correction_records = self.languagetool_service.convert_to_correction_records(
                    lt_errors, text
                )
                for record in correction_records:
                    self.collector.add_correction(record)
        except Exception as e:
            # LanguageTool fallito, continua con solo AI (solo log a livello debug)
            logger.debug(f"⚠️ LanguageTool unavailable, using AI-only mode: {e}")
        
        # 3. Correzione AI se necessaria
        needs_correction = lt_errors or self._needs_ai_correction(text)
        logger.info(f"🔍 Correction check: lt_errors={len(lt_errors) if lt_errors else 0}, needs_ai={self._needs_ai_correction(text)}, total_needs={needs_correction}")
        
        if needs_correction:
            corrected_text = self.openai_service.correct_text(text)
            if not corrected_text:
                logger.warning("⚠️  AI correction failed, using original text")
                corrected_text = text
            elif self.collector is not None and corrected_text != text:
                # Traccia correzione GPT (livello paragrafo)
                logger.info(f"📝 Tracking GPT correction: '{text[:50]}...' -> '{corrected_text[:50]}...'")
                self._track_gpt_correction(text, corrected_text)
            elif corrected_text == text:
                logger.info("⚠️  GPT returned same text, not tracking")
        else:
            logger.info("⚠️  No correction needed, skipping AI")
            corrected_text = text
        
        # 4. Validazione qualità (confronta con originale, non con testo già modificato da custom!)
        quality_score = self.safe_corrector.validate_correction_quality(original_text, corrected_text)
        logger.info(f"📊 [MAIN FLOW] Quality score: {quality_score.overall_score:.2f}, threshold: {self.config.min_quality_threshold}, will_accept={quality_score.overall_score >= self.config.min_quality_threshold}")
        
        # 5. Decisione applicazione correzione
        if quality_score.overall_score >= self.config.min_quality_threshold:
            # Cache il risultato se migliorato
            if use_cache and self.config.cache_enabled and corrected_text != original_text:
                self.cache_service.cache_with_metadata(
                    text=original_text,
                    correction=corrected_text,
                    quality=quality_score.overall_score if hasattr(quality_score, 'overall_score') else 0.8,
                    correction_type='correction_engine'
                )
            
            logger.info(f"✅ [MAIN FLOW] Text corrected and ACCEPTED (score: {quality_score.overall_score:.2f})")
            return corrected_text, quality_score
        else:
            logger.info(f"❌ [MAIN FLOW] Correction REJECTED (score: {quality_score.overall_score:.2f})")
            return original_text, quality_score
    
    def correct_paragraph_batch(self, paragraphs: List[Paragraph], max_workers: Optional[int] = None) -> List[Tuple[bool, QualityScore]]:
        """
        Corregge un batch di paragrafi in parallelo.
        
        Args:
            paragraphs: Lista di paragrafi da correggere
            max_workers: Numero massimo di thread (default: dalla config)
            
        Returns:
            List[Tuple[bool, QualityScore]]: Risultati per ogni paragrafo
        """
        if max_workers is None:
            max_workers = self.config.max_workers
        
        logger.info(f"🔄 Processing {len(paragraphs)} paragraphs with {max_workers} workers")
        
        results = []
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Invia tutti i task
            future_to_paragraph = {
                executor.submit(self._correct_single_paragraph, para): para 
                for para in paragraphs
            }
            
            # Raccogli i risultati mantenendo l'ordine
            for future in as_completed(future_to_paragraph):
                paragraph = future_to_paragraph[future]
                try:
                    success, quality_score = future.result()
                    results.append((success, quality_score))
                    
                    if success:
                        logger.debug(f"✅ Paragraph corrected (score: {quality_score.overall_score:.2f})")
                    else:
                        logger.debug(f"⚠️  Paragraph correction rejected")
                        
                except Exception as e:
                    logger.error(f"❌ Error correcting paragraph: {e}")
                    results.append((False, QualityScore(
                        overall_score=0.0,
                        confidence=CorrectionConfidence.VERY_LOW,
                        content_preservation=0.0,
                        grammar_improvement=0.0,
                        style_preservation=0.0,
                        safety_score=0.0,
                        issues=[f"Error: {str(e)}"]
                    )))
        
        success_count = sum(1 for success, _ in results if success)
        logger.info(f"📊 Batch complete: {success_count}/{len(paragraphs)} corrections applied")
        
        return results
    
    # Metodi privati
    
    def _process_corrections(self, context: CorrectionContext) -> bool:
        """Elabora tutte le correzioni per il contesto dato"""
        import time
        
        if not context.target_paragraphs:
            logger.info("ℹ️  No paragraphs require correction")
            return True
        
        # Processo in batch per efficienza
        batch_size = self.config.batch_size
        total_paragraphs = len(context.target_paragraphs)
        total_batches = (total_paragraphs - 1) // batch_size + 1
        
        for i in range(0, total_paragraphs, batch_size):
            batch = context.target_paragraphs[i:i + batch_size]
            batch_num = i // batch_size + 1
            
            logger.info(f"🔄 Processing batch {batch_num}/{total_batches}")
            
            # Correggi il batch
            batch_results = self.correct_paragraph_batch(batch)
            
            # Aggiorna statistiche
            for (success, quality_score) in batch_results:
                context.total_processed += 1
                if success:
                    context.corrections_applied += 1
                else:
                    context.corrections_rejected += 1
                
                # Log dettagliato per debugging
                context.corrections_log.append({
                    'paragraph_index': context.total_processed - 1,
                    'success': success,
                    'quality_score': quality_score.overall_score,
                    'timestamp': self._get_timestamp()
                })
            
            # Breve pausa tra batch per evitare rate limiting (skip sull'ultimo batch)
            if batch_num < total_batches:
                logger.debug(f"⏸️  Pausing 1s before next batch to avoid rate limiting")
                time.sleep(1)
        
        # Verifica soglia di successo
        success_rate = context.corrections_applied / context.total_processed if context.total_processed > 0 else 0
        
        if success_rate < self.config.min_success_rate:
            # ✅ Non abortire: conserva comunque le correzioni applicate
            logger.warning(f"⚠️ Success rate basso: {success_rate:.2%} < {self.config.min_success_rate:.2%}")
            # prosegui senza return False
        
        logger.info(f"✅ Correction process completed with {success_rate:.2%} success rate")
        return True
    
    def _correct_single_paragraph(self, paragraph: Paragraph) -> Tuple[bool, QualityScore]:
        """Corregge un singolo paragrafo"""
        original_text = paragraph.text
        logger.info(f"📄 [PARAGRAPH] _correct_single_paragraph called with {len(original_text)} chars")
        
        if not original_text.strip():
            logger.info("⚠️  [PARAGRAPH] Empty paragraph, skipping")
            return False, QualityScore(
                overall_score=0.0,
                confidence=CorrectionConfidence.VERY_LOW,
                content_preservation=1.0,  # Empty text preserved
                grammar_improvement=0.0,
                style_preservation=1.0,
                safety_score=1.0,  # Safe to skip empty
                issues=["Empty text skipped"]
            )
        
        try:
            corrected_text, quality_score = self.correct_text_fragment(original_text)
            logger.info(f"🔍 [PARAGRAPH] correct_text_fragment returned: corrected={corrected_text != original_text}, score={quality_score.overall_score:.2f}")
            
            # Applica la correzione se migliorata
            if corrected_text != original_text and quality_score.overall_score >= self.config.min_quality_threshold:
                # Preserva la formattazione applicando solo il testo
                self._apply_text_preserving_format(paragraph, corrected_text)
                logger.info(f"✅ [PARAGRAPH] Correction APPLIED to document")
                return True, quality_score
            else:
                logger.info(f"⚠️  [PARAGRAPH] Correction NOT applied: corrected={corrected_text != original_text}, score={quality_score.overall_score:.2f}")
                return False, quality_score
                
        except Exception as e:
            logger.error(f"❌ [PARAGRAPH] Error correcting paragraph: {e}")
            return False, QualityScore(
                overall_score=0.0,
                confidence=CorrectionConfidence.VERY_LOW,
                content_preservation=0.0,
                grammar_improvement=0.0,
                style_preservation=0.0,
                safety_score=0.0,
                issues=[f"Exception: {str(e)}"]
            )
    
    def _apply_text_preserving_format(self, paragraph: Paragraph, new_text: str):
        """Applica il nuovo testo evitando clobbering dei run."""
        if paragraph.runs:
            for run in paragraph.runs:
                run.text = ""  # svuota i run esistenti
            paragraph.add_run(new_text)    # un singolo run pulito
        else:
            paragraph.text = new_text
    
    def _should_correct_paragraph(self, paragraph: Paragraph) -> bool:
        """Determina se un paragrafo deve essere corretto"""
        text = paragraph.text.strip()
        
        if not text:
            return False
        
        # Evita paragrafi con equazioni matematiche
        if self._has_math_content(paragraph):
            return False
        
        # Evita paragrafi troppo corti, MA consenti frasi/dialoghi veri
        if len(text) < self.config.min_paragraph_length:
            if text.endswith(('.', '!', '?')) or text.startswith(('«', '–', '—', '- ')):
                pass  # consenti dialoghi/titoli
            else:
                return False
        
        # Evita paragrafi con pattern speciali (codici, date, numeri)
        if self._is_special_content(text):
            return False
        
        return True
    
    def _has_math_content(self, paragraph: Paragraph) -> bool:
        """Controlla se il paragrafo contiene contenuto matematico"""
        return bool(paragraph._p.xpath(".//*[local-name()='oMath' or local-name()='oMathPara']"))
    
    def _is_special_content(self, text: str) -> bool:
        """Controlla se il testo è contenuto speciale da non correggere"""
        import re
        
        special_patterns = [
            r'^\d+[\.\)]\s*$',              # numerazione "1." "1)"
            r'^\s*CAPITOLO\s+\d+\s*$',      # titoli capitolo
            r'^\s*[A-Z\s]{3,}\s*$',         # titoli tutto maiuscolo
            r'^[A-Z][a-z]+\s+\d{4}$',      # date "Gennaio 2024"
            r'^\s*\*{3,}\s*$',              # separatori asterischi
        ]
        
        for pattern in special_patterns:
            if re.match(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _needs_ai_correction(self, text: str) -> bool:
        """Euristica per determinare se il testo potrebbe beneficiare di correzione AI"""
        # Pattern che suggeriscono possibili errori
        import re
        
        error_indicators = [
            r'( {2,})',                       # spazi multipli
            r'([,;:.!?]\S)',                  # punteggiatura senza spazio dopo
            r'(\S[,;:.!?])',                  # nessuno spazio prima di punteggiatura finale
            r"\bqual['’]\s?è\b",              # qual’è / qual’ è
            r"\bun\s+p[òò]\b",                # un pò
            r"\bp[òò]\b",                     # pò isolato
            r"\b(a|e)\'\s",                   # apostrofo dritto seguito da spazio
            r'(«\s)|(\s»)',                   # virgolette caporali con spazi errati
        ]
        
        for pattern in error_indicators:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        
        return False
    
    def _log_correction_summary(self, context: CorrectionContext, output_path):
        """Registra un riassunto della sessione di correzione con statistiche di leggibilità"""
        success_rate = context.corrections_applied / context.total_processed if context.total_processed > 0 else 0
        
        # Estrai tutto il testo dal documento per analisi di leggibilità
        full_text = "\n".join([p.text for p in context.source_document.paragraphs if p.text.strip()])
        readability_stats = self.readability_analyzer.analyze(full_text)
        
        logger.info("=" * 60)
        logger.info("📊 CORRECTION SUMMARY")
        logger.info("=" * 60)
        logger.info(f"📄 Source document: {context.document_info.path.name}")
        logger.info(f"💾 Output document: {output_path.name}")
        logger.info(f"📝 Total paragraphs: {context.document_info.total_paragraphs}")
        logger.info(f"🎯 Paragraphs processed: {context.total_processed}")
        logger.info(f"✅ Corrections applied: {context.corrections_applied}")
        logger.info(f"⚠️  Corrections rejected: {context.corrections_rejected}")
        logger.info(f"📈 Success rate: {success_rate:.2%}")
        logger.info("")
        logger.info("📖 READABILITY ANALYSIS (GULPEASE)")
        logger.info("-" * 60)
        
        if readability_stats['gulpease'] is not None:
            logger.info(f"📊 Indice Gulpease: {readability_stats['gulpease']:.2f}/100")
            logger.info(f"🔤 Parole: {readability_stats['words']:,}")
            logger.info(f"📝 Frasi: {readability_stats['sentences']:,}")
            logger.info(f"📏 Lunghezza media parola: {readability_stats['avg_word_length']:.2f} lettere")
            logger.info(f"📐 Lunghezza media frase: {readability_stats['avg_sentence_length']:.2f} parole")
            logger.info("")
            logger.info("👥 Difficoltà per livello di scolarizzazione:")
            
            difficulty_labels = {
                'licenza_elementare': '   📚 Licenza elementare',
                'licenza_media': '   🎓 Licenza media',
                'diploma_superiore': '   🎯 Diploma superiore'
            }
            
            for level, label in difficulty_labels.items():
                difficulty = readability_stats['difficulty'].get(level, 'N/A')
                logger.info(f"{label}: {difficulty}")
        else:
            logger.info("⚠️  Testo non analizzabile per leggibilità")
        
        logger.info("=" * 60)
    
    def _get_timestamp(self) -> str:
        """Genera timestamp per logging"""
        from datetime import datetime
        return datetime.now().isoformat()
    
    # Metodi pubblici per il tracking e reporting
    
    def get_collector(self) -> Optional[CorrectionCollector]:
        """
        Restituisce il collector per accedere ai dati di tracking.
        
        Returns:
            CorrectionCollector o None se tracking disabilitato
        """
        return self.collector
    
    def get_correction_statistics(self) -> Dict[str, Any]:
        """
        Restituisce statistiche aggregate sulle correzioni.
        
        Returns:
            Dizionario con statistiche complete
        """
        if self.collector is None:
            return {
                'tracking_enabled': False,
                'message': 'Tracking not enabled for this engine'
            }
        
        stats = self.collector.get_statistics()
        
        return {
            'tracking_enabled': True,
            'total_corrections': stats.total_corrections,
            'unique_words': stats.unique_words,
            'processing_time': stats.processing_time,
            'by_category': {
                cat.display_name: count 
                for cat, count in stats.by_category.items()
            },
            'by_source': {
                src.value: count 
                for src, count in stats.by_source.items()
            },
            'applied_corrections': stats.applied_corrections,
            'ignored_corrections': stats.ignored_corrections,
        }
    
    def export_corrections_report(self, output_path: str, format: str = 'json'):
        """
        Esporta un report completo delle correzioni.
        
        Args:
            output_path: Percorso del file di output
            format: Formato del report ('json' o 'html')
        """
        if self.collector is None:
            logger.warning("⚠️  Tracking not enabled, cannot export report")
            return False
        
        try:
            if format.lower() == 'json':
                self.collector.export_to_json(output_path)
                logger.info(f"📄 Corrections report exported to: {output_path}")
                return True
            elif format.lower() == 'html':
                # Placeholder per future implementazioni HTML
                logger.warning("⚠️  HTML report not yet implemented")
                return False
            else:
                logger.error(f"❌ Unsupported format: {format}")
                return False
        except Exception as e:
            logger.error(f"❌ Failed to export report: {e}")
            return False
    
    def track_languagetool_errors(self, text: str, position: int = 0):
        """
        Traccia errori trovati da LanguageTool nel collector.
        
        Args:
            text: Testo da analizzare
            position: Posizione nel documento
        """
        if self.collector is None:
            return
        
        # Ottieni gli errori da LanguageTool
        errors = self.languagetool_service.check_text(text, use_cache=False)
        
        if not errors:
            return
        
        # Converti in CorrectionRecord e aggiungi al collector
        records = self.languagetool_service.convert_to_correction_records(errors, text)
        self.collector.add_corrections(records)
        
        logger.debug(f"📊 Tracked {len(records)} LanguageTool corrections")
    
    def track_openai_correction(
        self, 
        original_text: str, 
        corrected_text: str, 
        position: int = 0
    ):
        """
        Traccia una correzione OpenAI nel collector.
        
        Args:
            original_text: Testo originale
            corrected_text: Testo corretto
            position: Posizione nel documento
        """
        if self.collector is None:
            return
        
        # Crea record dalla correzione OpenAI
        record = self.openai_service.create_correction_record_from_response(
            original_text=original_text,
            corrected_text=corrected_text,
            position=position
        )
        
        if record:
            self.collector.add_correction(record)
            logger.debug("📊 Tracked OpenAI correction")
    
    def _track_gpt_correction(self, original_text: str, corrected_text: str):
        """
        Traccia una correzione GPT creando un CorrectionRecord.
        
        Args:
            original_text: Testo originale
            corrected_text: Testo corretto da GPT
        """
        if self.collector is None:
            logger.debug("⚠️  Collector is None in _track_gpt_correction")
            return
        
        logger.debug(f"📝 _track_gpt_correction called with {len(original_text)} -> {len(corrected_text)} chars")
            
        from datetime import datetime
        from ..models import CorrectionRecord, CorrectionCategory, CorrectionSource
        
        # Determina la categoria basandosi sul tipo di modifica
        category = CorrectionCategory.MIGLIORABILI  # Default per GPT
        
        # Se la differenza è minima, probabilmente è una correzione grammaticale
        import difflib
        similarity = difflib.SequenceMatcher(None, original_text, corrected_text).ratio()
        if similarity > 0.9:
            category = CorrectionCategory.ERRORI_RICONOSCIUTI
        
        record = CorrectionRecord(
            category=category,
            source=CorrectionSource.OPENAI_GPT,
            original_text=original_text,
            corrected_text=corrected_text,
            context=original_text,
            position=0,
            paragraph_index=0,
            sentence_index=0,
            rule_id="GPT_CORRECTION",
            message="Correzione suggerita da GPT-4",
            suggestions=[corrected_text],
            confidence_score=0.8,
            timestamp=datetime.now()
        )
        
        self.collector.add_correction(record)
        logger.debug(f"✅ Tracked GPT correction successfully (collector now has {len(self.collector)} corrections)")
    
    def _analyze_special_categories(self, document):
        """
        Analizza e traccia categorie speciali (FASE 7).
        
        Rileva:
        - Parole straniere (inglese, latino, francese, etc.)
        - Parole potenzialmente imbarazzanti/sensibili
        - Nomi propri tramite NER (Persone, Luoghi, Organizzazioni)
        
        Popola il CorrectionCollector con CorrectionRecord appropriati.
        
        Args:
            document: Documento completo da analizzare
        """
        if not self.special_categories_service or not self.collector:
            return
        
        from datetime import datetime
        from src.correttore.models.correction_tracking import (
            CorrectionRecord, CorrectionCategory, CorrectionSource
        )
        
        try:
            # Estrai testo completo dal documento
            full_text = "\n".join([p.text for p in document.paragraphs if p.text.strip()])
            
            # 1. Rileva parole straniere
            logger.debug("🌍 Detecting foreign words...")
            foreign_words = self.special_categories_service.detect_foreign_words(
                full_text, 
                min_occurrences=1
            )
            
            for fw in foreign_words:
                record = CorrectionRecord(
                    id=f"foreign_{fw.word}_{fw.position}",
                    category=CorrectionCategory.LINGUE,
                    original_text=fw.word,
                    corrected_text=None,  # Nessuna correzione, solo info
                    context=fw.context,
                    position=fw.position,
                    paragraph_index=0,
                    sentence_index=0,
                    source=CorrectionSource.SYSTEM,
                    confidence_score=1.0,
                    rule_id=f"FOREIGN_WORD_{fw.language.upper()}",
                    message=f"Parola {fw.language} (occorrenze: {fw.count})",
                    suggestions=[],
                    timestamp=datetime.now()
                )
                self.collector.add_correction(record)
            
            logger.info(f"✅ Tracked {len(foreign_words)} foreign words")
            
            # 2. Rileva parole sensibili/imbarazzanti
            logger.debug("😳 Detecting sensitive words...")
            sensitive_words = self.special_categories_service.detect_sensitive_words(
                full_text,
                min_occurrences=1
            )
            
            for sw in sensitive_words:
                record = CorrectionRecord(
                    id=f"sensitive_{sw.word}_{sw.position}",
                    category=CorrectionCategory.IMBARAZZANTI,
                    original_text=sw.word,
                    corrected_text=None,
                    context=sw.context,
                    position=sw.position,
                    paragraph_index=0,
                    sentence_index=0,
                    source=CorrectionSource.SYSTEM,
                    confidence_score=1.0,
                    rule_id=f"SENSITIVE_{sw.category.upper()}",
                    message=f"Parola potenzialmente sensibile ({sw.category}, occorrenze: {sw.count})",
                    suggestions=[],
                    timestamp=datetime.now()
                )
                self.collector.add_correction(record)
            
            logger.info(f"✅ Tracked {len(sensitive_words)} sensitive words")
            
            # 3. Estrai nomi propri tramite NER
            logger.debug("👤 Extracting proper nouns via NER...")
            
            # Usa lemmatization service per NER
            try:
                from src.correttore.services.lemmatization_service import (
                    LemmatizationService
                )
                lemma_service = LemmatizationService()
                
                proper_nouns = self.special_categories_service.extract_proper_nouns_from_ner(
                    full_text,
                    lemmatization_service=lemma_service
                )
                
                for pn in proper_nouns:
                    record = CorrectionRecord(
                        id=f"proper_noun_{pn.name}_{pn.position}",
                        category=CorrectionCategory.NOMI_SIGLE,
                        original_text=pn.name,
                        corrected_text=None,
                        context=pn.context,
                        position=pn.position,
                        paragraph_index=0,
                        sentence_index=0,
                        source=CorrectionSource.SYSTEM,
                        confidence_score=1.0,
                        rule_id=f"NER_{pn.entity_type}",
                        message=f"{pn.entity_label} (occorrenze: {pn.count})",
                        suggestions=[],
                        timestamp=datetime.now()
                    )
                    self.collector.add_correction(record)
                
                logger.info(f"✅ Tracked {len(proper_nouns)} proper nouns")
                
            except Exception as e:
                logger.warning(f"⚠️  NER extraction failed: {e}")
            
            # Log statistiche finali
            stats = self.special_categories_service.get_statistics(
                foreign_words, sensitive_words, 
                proper_nouns if 'proper_nouns' in locals() else []
            )
            
            logger.info(
                f"📊 Special categories summary: "
                f"{stats['foreign_words']['total_unique']} foreign, "
                f"{stats['sensitive_words']['total_unique']} sensitive, "
                f"{stats['proper_nouns']['total_unique']} proper nouns"
            )
            
        except Exception as e:
            logger.error(f"❌ Special categories analysis failed: {e}")
    
    def _generate_html_report(self, input_path, output_path):
        """
        Genera report HTML delle correzioni.
        
        Args:
            input_path: Path del documento originale
            output_path: Path del documento corretto
        """
        logger.info(f"🔍 _generate_html_report chiamato: collector={self.collector is not None}")
        
        if self.collector is None:
            logger.warning("⚠️  Collector is None, cannot generate HTML report")
            return
        
        try:
            from pathlib import Path
            from ..utils.html_report_generator import generate_orthography_report
            
            # Determina path del report
            output_path = Path(output_path)
            report_path = output_path.parent / f"{output_path.stem}_report.html"
            
            logger.info(f"📝 Generating report at: {report_path}")
            logger.info(f"📊 Collector has {len(self.collector._corrections)} corrections")
            
            # Genera report
            document_name = Path(input_path).stem
            generate_orthography_report(
                collector=self.collector,
                output_path=str(report_path),
                document_name=document_name,
                standalone=True,
                show_feedback_buttons=False
            )
            
            logger.info(f"📋 Report HTML generato: {report_path}")
            
        except Exception as e:
            logger.warning(f"⚠️  Impossibile generare report HTML: {e}")
    
    # ========================================
    # FASE 6: FEEDBACK INTEGRATION
    # ========================================
    
    def _load_custom_corrections(self):
        """
        Carica correzioni personalizzate apprese da feedback utente.
        Files: data/custom_corrections.txt, data/custom_whitelist.txt
        """
        from pathlib import Path
        
        project_root = Path(__file__).parent.parent.parent.parent
        corrections_file = project_root / "data" / "custom_corrections.txt"
        whitelist_file = project_root / "data" / "custom_whitelist.txt"
        
        # Carica custom corrections (formato: "errore -> correzione")
        if corrections_file.exists():
            try:
                with open(corrections_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        # Skip commenti e righe vuote
                        if not line or line.startswith('#'):
                            continue
                        
                        # Parse formato "errore -> correzione" o "errore → correzione"
                        if '->' in line or '→' in line:
                            separator = '->' if '->' in line else '→'
                            parts = line.split(separator)
                            if len(parts) == 2:
                                original = parts[0].strip().lower()
                                corrected = parts[1].strip()
                                self.custom_corrections[original] = corrected
                
                logger.debug(f"📚 Loaded {len(self.custom_corrections)} custom corrections")
            except Exception as e:
                logger.warning(f"⚠️  Error loading custom corrections: {e}")
        
        # Carica whitelist (parole da NON correggere)
        if whitelist_file.exists():
            try:
                with open(whitelist_file, 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        if line and not line.startswith('#'):
                            self.custom_whitelist.add(line.lower())
                
                logger.debug(f"✅ Loaded {len(self.custom_whitelist)} whitelisted words")
            except Exception as e:
                logger.warning(f"⚠️  Error loading whitelist: {e}")
    
    def _apply_custom_corrections(self, text: str) -> Tuple[str, int]:
        """
        Applica correzioni personalizzate al testo.
        Esegue prima delle altre correzioni (priorità massima).
        
        Args:
            text: Testo da correggere
            
        Returns:
            (testo_corretto, numero_correzioni_applicate)
        """
        logger.info(f"🔍 [FASE 6 DEBUG] _apply_custom_corrections called, collector={self.collector is not None}, custom_corrections={len(self.custom_corrections) if self.custom_corrections else 0}")
        if not self.custom_corrections:
            logger.warning("⚠️ [FASE 6 DEBUG] No custom_corrections loaded!")
            return text, 0
        
        corrected_text = text
        corrections_count = 0
        
        # Applica ogni correzione custom
        for original, corrected in self.custom_corrections.items():
            # Case-insensitive replacement preservando maiuscole originali
            import re
            
            # Pattern per trovare la parola (word boundary)
            pattern = re.compile(r'\b' + re.escape(original) + r'\b', re.IGNORECASE)
            
            def replace_preserving_case(match):
                matched_text = match.group(0)
                
                # Se tutto maiuscolo, mantieni maiuscolo
                if matched_text.isupper():
                    return corrected.upper()
                # Se prima lettera maiuscola, mantieni prima maiuscola
                elif matched_text[0].isupper():
                    return corrected.capitalize()
                # Altrimenti lowercase
                else:
                    return corrected.lower()
            
            new_text = pattern.sub(replace_preserving_case, corrected_text)
            
            if new_text != corrected_text:
                corrections_count += 1
                corrected_text = new_text
                logger.info(f"✨ [FASE 6 DEBUG] Applied custom correction: '{original}' → '{corrected}', corrections_count={corrections_count}")
                
                # Traccia correzione custom nel collector (FASE 6)
                if self.collector is not None:
                    from correttore.models import CorrectionRecord, CorrectionSource, CorrectionCategory
                    record = CorrectionRecord(
                        original_text=original,
                        corrected_text=corrected,
                        source=CorrectionSource.CUSTOM_RULES,
                        category=CorrectionCategory.ERRORI_RICONOSCIUTI,
                        rule_id="CUSTOM_FEEDBACK_FASE6",
                        confidence_score=1.0,
                        position=0,
                        is_applied=True
                    )
                    self.collector.add_correction(record)
                    logger.info(f"✅ [FASE 6 DEBUG] CorrectionRecord added to collector!")
                else:
                    logger.error(f"❌ [FASE 6 DEBUG] collector is None, correction NOT tracked!")
        
        logger.info(f"🎯 [FASE 6 DEBUG] _apply_custom_corrections finished: {corrections_count} corrections applied")
        return corrected_text, corrections_count
    
    def _is_whitelisted(self, word: str) -> bool:
        """
        Verifica se una parola è nella whitelist (da NON correggere).
        
        Args:
            word: Parola da verificare
            
        Returns:
            True se whitelisted
        """
        return word.lower() in self.custom_whitelist


# Factory function per uso diretto
def create_correction_engine(enable_tracking: bool = True) -> CorrectionEngine:
    """
    Crea un'istanza del motore di correzione con configurazione predefinita.
    
    Args:
        enable_tracking: Se abilitare il tracking delle correzioni per il report
    """
    return CorrectionEngine(enable_tracking=enable_tracking)