# core/formatting_manager.py
"""
Gestione avanzata della formattazione nei documenti Word.
Preserva stili, corsivi, grassetti e formattazione complessa durante le correzioni.
"""

import logging
import re
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass
from copy import deepcopy

from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.shared import qn
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)

@dataclass
class RunFormatting:
    """Formattazione di un singolo run"""
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    font_name: Optional[str] = None
    font_size: Optional[int] = None
    font_color: Optional[str] = None
    
    def __eq__(self, other):
        if not isinstance(other, RunFormatting):
            return False
        return (self.bold == other.bold and 
                self.italic == other.italic and
                self.underline == other.underline and
                self.font_name == other.font_name and
                self.font_size == other.font_size and
                self.font_color == other.font_color)

@dataclass
class FormattedSegment:
    """Segmento di testo con la sua formattazione"""
    text: str
    formatting: RunFormatting
    start_pos: int
    end_pos: int

@dataclass
class FormattingMap:
    """Mappa della formattazione di un paragrafo"""
    segments: List[FormattedSegment]
    original_text: str
    paragraph_style: Optional[str] = None

class FormattingManager:
    """
    Gestore della formattazione per preservare stili durante le correzioni.
    Mantiene mappature dettagliate di corsivi, grassetti e altri stili.
    """
    
    def __init__(self):
        self.format_cache: Dict[str, FormattingMap] = {}
    
    def extract_formatting(self, paragraph: Paragraph) -> FormattingMap:
        """
        Estrae la formattazione completa di un paragrafo.
        
        Args:
            paragraph: Paragrafo da analizzare
            
        Returns:
            FormattingMap: Mappa completa della formattazione
        """
        segments = []
        current_pos = 0
        original_text = paragraph.text
        
        for run in paragraph.runs:
            if not run.text:
                continue
                
            run_formatting = self._extract_run_formatting(run)
            
            segment = FormattedSegment(
                text=run.text,
                formatting=run_formatting,
                start_pos=current_pos,
                end_pos=current_pos + len(run.text)
            )
            segments.append(segment)
            current_pos += len(run.text)
        
        paragraph_style = paragraph.style.name if paragraph.style else None
        
        formatting_map = FormattingMap(
            segments=segments,
            original_text=original_text,
            paragraph_style=paragraph_style
        )
        
        # Cache per riutilizzo
        cache_key = f"{id(paragraph)}_{hash(original_text)}"
        self.format_cache[cache_key] = formatting_map
        
        logger.debug(f"📝 Extracted formatting for paragraph: {len(segments)} segments")
        
        return formatting_map
    
    def _extract_run_formatting(self, run: Run) -> RunFormatting:
        """Estrae formattazione di un singolo run"""
        formatting = RunFormatting()
        
        # Font properties
        if run.font.bold is not None:
            formatting.bold = run.font.bold
        if run.font.italic is not None:
            formatting.italic = run.font.italic
        if run.font.underline is not None:
            formatting.underline = bool(run.font.underline)
        
        # Font name
        if run.font.name:
            formatting.font_name = run.font.name
        
        # Font size
        if run.font.size:
            formatting.font_size = run.font.size.pt
        
        # Font color
        if run.font.color.rgb:
            formatting.font_color = str(run.font.color.rgb)
        
        return formatting
    
    def apply_formatting(self, paragraph: Paragraph, corrected_text: str, 
                        original_formatting: FormattingMap) -> bool:
        """
        Applica la formattazione originale al testo corretto.
        
        Args:
            paragraph: Paragrafo da modificare
            corrected_text: Testo corretto
            original_formatting: Formattazione originale
            
        Returns:
            bool: True se la formattazione è stata applicata con successo
        """
        try:
            # Pulisce i run esistenti
            for run in paragraph.runs:
                run._element.getparent().remove(run._element)
            
            # Mappa il testo corretto alla formattazione originale
            mapped_segments = self._map_formatting_to_corrected_text(
                corrected_text, original_formatting
            )
            
            # Crea nuovi run con formattazione
            for segment in mapped_segments:
                if segment.text:
                    new_run = paragraph.add_run(segment.text)
                    self._apply_run_formatting(new_run, segment.formatting)
            
            logger.debug(f"✅ Applied formatting to corrected text: {len(mapped_segments)} segments")
            return True
            
        except Exception as e:
            logger.error(f"❌ Failed to apply formatting: {e}")
            # Fallback: usa testo senza formattazione
            paragraph.clear()
            paragraph.add_run(corrected_text)
            return False
    
    def _map_formatting_to_corrected_text(self, corrected_text: str, 
                                        original_formatting: FormattingMap) -> List[FormattedSegment]:
        """
        Mappa la formattazione originale al testo corretto usando fuzzy matching.
        Preserva formattazione anche quando il testo è stato modificato.
        """
        if not original_formatting.segments:
            # Nessuna formattazione originale
            return [FormattedSegment(
                text=corrected_text,
                formatting=RunFormatting(),
                start_pos=0,
                end_pos=len(corrected_text)
            )]
        
        mapped_segments = []
        corrected_words = self._tokenize_for_formatting(corrected_text)
        original_segments_words = []
        
        # Tokenizza ogni segmento originale
        for segment in original_formatting.segments:
            words = self._tokenize_for_formatting(segment.text)
            original_segments_words.append((segment, words))
        
        # Mappa parole corrette ai segmenti originali
        corrected_pos = 0
        
        for corrected_word in corrected_words:
            best_match = self._find_best_format_match(corrected_word, original_segments_words)
            
            if best_match:
                segment, formatting = best_match
                mapped_segment = FormattedSegment(
                    text=corrected_word,
                    formatting=formatting,
                    start_pos=corrected_pos,
                    end_pos=corrected_pos + len(corrected_word)
                )
                mapped_segments.append(mapped_segment)
            else:
                # Usa formattazione neutra per parole non mappabili
                mapped_segment = FormattedSegment(
                    text=corrected_word,
                    formatting=RunFormatting(),
                    start_pos=corrected_pos,
                    end_pos=corrected_pos + len(corrected_word)
                )
                mapped_segments.append(mapped_segment)
            
            corrected_pos += len(corrected_word)
        
        # Consolida segmenti consecutivi con stessa formattazione
        return self._consolidate_segments(mapped_segments)
    
    def _tokenize_for_formatting(self, text: str) -> List[str]:
        """
        Tokenizza testo preservando spazi e punteggiatura per formattazione.
        """
        # Pattern che cattura parole, spazi e punteggiatura separatamente
        tokens = re.findall(r'\S+|\s+', text)
        return tokens
    
    def _find_best_format_match(self, word: str, segments_words: List[Tuple]) -> Optional[Tuple]:
        """
        Trova il miglior match di formattazione per una parola.
        Usa similarity matching per gestire correzioni.
        """
        word_lower = word.lower().strip()
        if not word_lower:
            return None
        
        best_score = 0.0
        best_match = None
        
        for segment, words in segments_words:
            for orig_word in words:
                orig_word_lower = orig_word.lower().strip()
                if not orig_word_lower:
                    continue
                
                # Exact match
                if word_lower == orig_word_lower:
                    return (segment, segment.formatting)
                
                # Similarity match (per correzioni)
                if len(word_lower) > 2 and len(orig_word_lower) > 2:
                    # Calcola similarity
                    common_chars = set(word_lower) & set(orig_word_lower)
                    similarity = len(common_chars) / max(len(word_lower), len(orig_word_lower))
                    
                    # Bonus per prefissi/suffissi comuni
                    if word_lower.startswith(orig_word_lower[:3]) or orig_word_lower.startswith(word_lower[:3]):
                        similarity += 0.2
                    if word_lower.endswith(orig_word_lower[-3:]) or orig_word_lower.endswith(word_lower[-3:]):
                        similarity += 0.2
                    
                    if similarity > best_score and similarity > 0.6:
                        best_score = similarity
                        best_match = (segment, segment.formatting)
        
        return best_match
    
    def _consolidate_segments(self, segments: List[FormattedSegment]) -> List[FormattedSegment]:
        """
        Consolida segmenti consecutivi con formattazione identica.
        Ottimizza il numero di run nel documento.
        """
        if not segments:
            return segments
        
        consolidated = []
        current_segment = segments[0]
        
        for next_segment in segments[1:]:
            if current_segment.formatting == next_segment.formatting:
                # Unisci i segmenti
                current_segment = FormattedSegment(
                    text=current_segment.text + next_segment.text,
                    formatting=current_segment.formatting,
                    start_pos=current_segment.start_pos,
                    end_pos=next_segment.end_pos
                )
            else:
                consolidated.append(current_segment)
                current_segment = next_segment
        
        consolidated.append(current_segment)
        return consolidated
    
    def _apply_run_formatting(self, run: Run, formatting: RunFormatting):
        """Applica formattazione a un run"""
        if formatting.bold is not None:
            run.font.bold = formatting.bold
        if formatting.italic is not None:
            run.font.italic = formatting.italic
        if formatting.underline is not None:
            run.font.underline = formatting.underline
        if formatting.font_name:
            run.font.name = formatting.font_name
        if formatting.font_size:
            run.font.size = Pt(formatting.font_size)
        if formatting.font_color:
            try:
                # Converte stringa colore in RGBColor
                color_hex = formatting.font_color.replace('#', '')
                if len(color_hex) == 6:
                    r = int(color_hex[0:2], 16)
                    g = int(color_hex[2:4], 16)
                    b = int(color_hex[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
            except (ValueError, AttributeError):
                # Ignora errori di colore
                pass
    
    def preserve_formatting_in_correction(self, paragraph: Paragraph, 
                                        original_text: str, 
                                        corrected_text: str) -> bool:
        """
        Preserva formattazione durante correzione di un paragrafo.
        Metodo principale da usare durante le correzioni.
        
        Args:
            paragraph: Paragrafo da correggere
            original_text: Testo originale
            corrected_text: Testo corretto
            
        Returns:
            bool: True se formattazione preservata con successo
        """
        if original_text == corrected_text:
            return True  # Nessuna modifica necessaria
        
        try:
            # 1. Estrai formattazione originale
            original_formatting = self.extract_formatting(paragraph)
            
            # 2. Applica formattazione al testo corretto
            success = self.apply_formatting(paragraph, corrected_text, original_formatting)
            
            if success:
                logger.debug("✅ Formatting preserved successfully")
            else:
                logger.warning("⚠️ Formatting preservation failed, using plain text")
            
            return success
            
        except Exception as e:
            logger.error(f"❌ Error preserving formatting: {e}")
            # Fallback sicuro: usa testo corretto senza formattazione
            paragraph.clear()
            paragraph.add_run(corrected_text)
            return False
    
    def validate_formatting_preservation(self, original_paragraph: Paragraph, 
                                       corrected_paragraph: Paragraph) -> Dict[str, Any]:
        """
        Valida che la formattazione sia stata preservata correttamente.
        
        Returns:
            Dict con metriche di validazione
        """
        try:
            original_fmt = self.extract_formatting(original_paragraph)
            corrected_fmt = self.extract_formatting(corrected_paragraph)
            
            # Conta formattazioni preservate
            original_formats = set()
            corrected_formats = set()
            
            for segment in original_fmt.segments:
                fmt_key = (segment.formatting.bold, segment.formatting.italic, 
                          segment.formatting.underline, segment.formatting.font_name)
                original_formats.add(fmt_key)
            
            for segment in corrected_fmt.segments:
                fmt_key = (segment.formatting.bold, segment.formatting.italic,
                          segment.formatting.underline, segment.formatting.font_name)
                corrected_formats.add(fmt_key)
            
            preservation_ratio = len(original_formats & corrected_formats) / max(len(original_formats), 1)
            
            return {
                "preservation_ratio": preservation_ratio,
                "original_format_count": len(original_formats),
                "corrected_format_count": len(corrected_formats),
                "preserved_formats": len(original_formats & corrected_formats),
                "success": preservation_ratio >= 0.8
            }
            
        except Exception as e:
            logger.error(f"❌ Error validating formatting: {e}")
            return {
                "preservation_ratio": 0.0,
                "success": False,
                "error": str(e)
            }
    
    def get_formatting_stats(self) -> Dict[str, Any]:
        """Ottieni statistiche sulla gestione formattazione"""
        return {
            "cached_formats": len(self.format_cache),
            "cache_keys": list(self.format_cache.keys())
        }
    
    def clear_cache(self):
        """Pulisce la cache delle formattazioni"""
        self.format_cache.clear()
        logger.debug("🧹 Formatting cache cleared")

def create_formatting_manager() -> FormattingManager:
    """Factory function per creare FormattingManager"""
    return FormattingManager()
