#!/usr/bin/env python3
"""
Test di regressione per validare la nuova pipeline invertita.
Crea un documento con i casi problematici e verifica le correzioni.
"""

import tempfile
import sys
from pathlib import Path
from docx import Document

def create_comprehensive_test_document():
    """Crea un documento di test completo con tutti i casi problematici"""
    doc = Document()
    
    # Aggiungi un titolo
    title = doc.add_heading('Test Documento - Pipeline Invertita', 0)
    
    # Casi problematici identificati nel piano
    problematic_cases = [
        # Casi principali che dovremmo risolvere
        "C'era una vlta un principe molto coraggioso che viveva nel castello.",
        "U giorno il principe decise di esplorare il bosco incantato.",
        "Visitò la bottaga del vecchio falegname nel villaggio.",
        "Il falegname alfredo era famoso per i suoi mobili di legno.",
        "ansiano nonno raccontava sempre storie di draghi e cavalieri.",
        "La nonna preparava un delizioso sugu con pomodori freschi.",
        
        # Casi che dovrebbero rimanere corretti (regressione test)
        "CAPTIOLO primo: L'inizio dell'avventura incredibile.",
        "I bambini carezzzzavano dolcemente il gattino trovato.",
        "Ho fatto un errore ma go risolto il problema rapidamente.",
        "Il fato ha voluto che ci incontrassimo proprio oggi.",
        
        # Casi misti per test completo
        "ansiano principe vlta visitò la bottaga dove alfredo lavorava.",
        "U giorno CAPTIOLO sugu carezzzzavano go fato insieme.",
    ]
    
    # Aggiungi i paragrafi di test
    doc.add_heading('Casi di Test Principali', level=1)
    for i, case in enumerate(problematic_cases, 1):
        p = doc.add_paragraph(f"{i}. {case}")
    
    # Aggiungi una sezione con testo normale per context
    doc.add_heading('Contesto Narrativo', level=1)
    context_text = """
    Questo è un documento di prova per verificare il funzionamento della nuova pipeline di correzione.
    Il sistema dovrebbe essere in grado di correggere errori ortografici e grammaticali mantenendo
    il significato originale del testo e preservando lo stile dell'autore.
    """
    doc.add_paragraph(context_text)
    
    return doc

def main():
    """Test principale"""
    print("🧪 Test di Regressione - Pipeline Invertita")
    print("="*60)
    
    # Crea documento di test
    test_doc = create_comprehensive_test_document()
    
    # Crea file temporaneo in una directory persistente per il test
    test_dir = Path("test_output")
    test_dir.mkdir(exist_ok=True)
    
    original_path = test_dir / "test_regressione_originale.docx"
    corrected_path = test_dir / "test_regressione_corretto.docx"
    
    # Salva il documento originale
    test_doc.save(str(original_path))
    print(f"✅ Documento di test creato: {original_path}")
    
    # Mostra contenuto per verifica
    print("\n📄 Contenuto del test:")
    for i, p in enumerate(test_doc.paragraphs, 1):
        if p.text.strip() and not p.text.startswith('Test Documento') and not p.text.startswith('Casi di Test') and not p.text.startswith('Contesto'):
            print(f"   {p.text}")
    
    print(f"\n🎯 Aspettative per la nuova pipeline:")
    print("="*60)
    print("CORREZIONI ATTESE:")
    print("   • vlta → volta (NON più → alta)")
    print("   • bottaga → bottega (NON più → bottaia)")
    print("   • sugu → sugo (NON più → suga)")
    print("   • U giorno → Un giorno")
    print("   • alfredo → Alfredo")
    print("   • ansiano → anziano")
    print("\nCORREZIONI DA PRESERVARE:")
    print("   • CAPTIOLO → CAPITOLO")
    print("   • carezzzzavano → carezzavano")
    print("   • go → ho")
    print("   • fato → fatto")
    
    print(f"\n📁 Files generati:")
    print(f"   • Originale: {original_path}")
    print(f"   • Da processare → {corrected_path}")
    
    print(f"\n🚀 Per eseguire il test completo:")
    print(f"   python main.py \"{original_path}\"")
    print(f"   # oppure")
    print(f"   python src/core/correttore.py")
    
    return original_path, corrected_path

if __name__ == "__main__":
    original, corrected = main()
    print(f"\n✅ Setup test completato!")
    print(f"Files pronti in: {original.parent}")
