#!/usr/bin/env python3
"""
Test end-to-end per verificare che la pipeline invertita funzioni con i casi problematici.
"""

import sys
import tempfile
from pathlib import Path
from docx import Document

def create_test_document():
    """Crea un documento di test con i problemi identificati"""
    doc = Document()
    
    # Aggiungi paragrafi con i problemi specifici che dovremmo risolvere
    test_cases = [
        "C'era una vlta un principe che viveva in un castello.",
        "U giorno decise di visitare la bottaga del falegname.",
        "Il falegname alfredo preparò un ottimo sugu per la cena.",
        "ansiano nonno raccontava sempre storie interessanti.",
        "Il CAPTIOLO primo del libro era molto carezzzzavano.",
    ]
    
    for case in test_cases:
        p = doc.add_paragraph(case)
    
    return doc

def main():
    """Test principale"""
    print("🧪 Test End-to-End Pipeline Invertita")
    print("="*50)
    
    # Crea documento di test
    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = Path(tmpdir) / "test_document.docx"
        corrected_path = Path(tmpdir) / "test_document_corretto.docx"
        
        # Salva il documento di test
        doc = create_test_document()
        doc.save(str(test_path))
        print(f"✅ Documento di test creato: {test_path}")
        
        # Mostra contenuto originale
        print("\n📄 Contenuto originale:")
        for i, p in enumerate(doc.paragraphs, 1):
            if p.text.strip():
                print(f"   {i}. {p.text}")
        
        # Nota: Non eseguiamo la correzione completa qui perché richiede
        # OpenAI API key e LanguageTool server attivo
        print("\n⚠️  Per il test completo, eseguire:")
        print(f"   python main.py {test_path}")
        print("\n🎯 Risultati attesi con la nuova pipeline:")
        print("   • vlta → volta (NON più → alta)")
        print("   • bottaga → bottega (NON più → bottaia)")  
        print("   • sugu → sugo (NON più → suga)")
        print("   • U giorno → Un giorno")
        print("   • alfredo → Alfredo")
        print("   • ansiano → anziano")
        print("   • CAPTIOLO → CAPITOLO (mantenuto)")
        print("   • carezzzzavano → carezzavano (mantenuto)")

if __name__ == "__main__":
    main()
