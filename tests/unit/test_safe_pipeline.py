# test_safe_pipeline.py
"""
Test della nuova pipeline di correzione sicura.
Verifica che il sistema integrato funzioni correttamente.
"""

import asyncio
import logging
from pathlib import Path
from docx import Document
from docx.text.paragraph import Paragraph

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(name)s: %(message)s')

# Importa i moduli del sistema
from src.core.safe_correction import SafeCorrector
from src.core.validation import DocumentValidator
from src.core.correttore import Modification, NAME_RE, GLOSSARY_STOP

def create_test_document():
    """Crea un documento di test con vari tipi di errori"""
    doc = Document()
    
    # Aggiungi paragrafi con errori diversi
    test_paragraphs = [
        "Questo è un paragrafo senza errori che non dovrebbe essere modificato.",
        "Questo paragrafo ha alcuni errori di ortographia che dovrebbero essere corretti.",
        "Un'altra frase con refuso e qualche problema grammaticale da sistemare.",
        "Questo testo ha pò errori invece di po' e altri problemi.",
        "",  # Paragrafo vuoto
        "CAPPITOLO 1",  # Errore in maiuscolo
        "La seguente frase ha molti errori ortographici e grammaticali che devono essere sistemati correttamente."
    ]
    
    for text in test_paragraphs:
        para = doc.add_paragraph(text)
    
    return doc

async def test_safe_pipeline():
    """Test della pipeline di correzione sicura completa"""
    print("\n=== TEST PIPELINE CORREZIONE SICURA ===")
    
    # Setup
    safe_corrector = SafeCorrector(conservative_mode=True, quality_threshold=0.85)
    doc = create_test_document()
    
    # Simula il glossario e altre strutture
    glossary = set()
    mods = []
    
    # Lista paragrafi 
    paragraphs = list(doc.paragraphs)
    print(f"📄 Created test document with {len(paragraphs)} paragraphs")
    
    # Mock client per test (non useremo davvero OpenAI)
    class MockAsyncClient:
        async def chat_completions_create(self, *args, **kwargs):
            return None
    
    # Test del spellcheck sicuro
    print("\n--- Test Spellcheck Sicuro ---")
    spellcheck_applied = 0
    
    for i, para in enumerate(paragraphs):
        if not para.text.strip():
            continue
            
        original = para.text
        
        # Mock spellcheck function
        def mock_spellcheck(text):
            corrections = {
                "ortographia": "ortografia",
                "refuso": "errore", 
                "pò": "po'",
                "CAPPITOLO": "CAPITOLO"
            }
            corrected = text
            for wrong, right in corrections.items():
                corrected = corrected.replace(wrong, right)
            return corrected
        
        # Test con safe corrector
        result = safe_corrector.correct_with_rollback(para, mock_spellcheck, "spellcheck")
        
        if result.applied:
            spellcheck_applied += 1
            print(f"✅ Paragraph {i+1}: Spellcheck applied (quality: {result.quality_score.overall_score:.1%})")
            print(f"   '{original[:50]}...' → '{result.corrected_text[:50]}...'")
        else:
            print(f"⏭️  Paragraph {i+1}: No spellcheck needed or rolled back")
            
    print(f"📊 Spellcheck: {spellcheck_applied}/{len([p for p in paragraphs if p.text.strip()])} applied")
    
    # Test del grammar check (simulato)
    print("\n--- Test Grammar Check Sicuro ---")
    grammar_applied = 0
    
    for i, para in enumerate(paragraphs):
        if not para.text.strip():
            continue
            
        original = para.text
        
        # Mock grammar check
        def mock_grammar(text):
            # Simula correzioni grammaticali semplici
            if "un'altra" in text.lower():
                return text.replace("un'altra", "un altro")
            if text.endswith(" qui."):
                return text.replace(" qui.", ".")
            return text
        
        corrected = mock_grammar(para.text)
        
        # Validazione manuale (come nel sistema reale)
        if corrected != original:
            from src.core.validation import validate_correction
            if validate_correction(original, corrected):
                para.text = corrected
                grammar_applied += 1
                print(f"✅ Paragraph {i+1}: Grammar applied")
                print(f"   '{original[:50]}...' → '{corrected[:50]}...'")
                
                # Simula aggiornamento glossario
                for name in NAME_RE.findall(corrected):
                    if name.upper() not in GLOSSARY_STOP:
                        glossary.add(name)
            else:
                print(f"🔄 Paragraph {i+1}: Grammar rolled back (validation failed)")
        else:
            print(f"⏭️  Paragraph {i+1}: No grammar changes needed")
            
    print(f"📊 Grammar: {grammar_applied}/{len([p for p in paragraphs if p.text.strip()])} applied")
    
    # Test AI correction (simulato)
    print("\n--- Test AI Correction ---")
    
    # Simula batch AI processing
    ai_texts = []
    for para in paragraphs:
        if para.text.strip():
            ai_texts.append(para.text)
    
    # Mock AI correction
    def mock_ai_batch(texts):
        corrected = []
        for text in texts:
            # Simula miglioramenti AI minimali
            if "problema grammaticale" in text:
                corrected.append(text.replace("problema grammaticale", "errore grammaticale"))
            elif "sistemati correttamente" in text:
                corrected.append(text.replace("sistemati correttamente", "corretti"))
            else:
                corrected.append(text)  # Nessun cambiamento
        return corrected
    
    ai_corrected_list = mock_ai_batch(ai_texts)
    ai_applied = 0
    
    para_idx = 0
    for i, para in enumerate(paragraphs):
        if not para.text.strip():
            continue
            
        original = para.text
        ai_corrected = ai_corrected_list[para_idx]
        para_idx += 1
        
        if ai_corrected != original:
            from src.core.validation import validate_correction
            if validate_correction(original, ai_corrected):
                para.text = ai_corrected
                ai_applied += 1
                
                # Simula registrazione modifica
                mods.append(Modification(i+1, original, ai_corrected))
                
                print(f"✅ Paragraph {i+1}: AI correction applied")
                print(f"   '{original[:50]}...' → '{ai_corrected[:50]}...'")
            else:
                print(f"🔄 Paragraph {i+1}: AI correction rolled back")
        else:
            print(f"⏭️  Paragraph {i+1}: No AI changes")
    
    print(f"📊 AI Correction: {ai_applied}/{len(ai_texts)} applied")
    
    # Statistiche finali
    total_corrections = spellcheck_applied + grammar_applied + ai_applied
    total_paragraphs = len([p for p in paragraphs if p.text.strip()])
    
    print(f"\n📊 STATISTICHE FINALI:")
    print(f"   • Paragrafi processati: {total_paragraphs}")
    print(f"   • Correzioni applicate: {total_corrections}")
    print(f"   • Modifiche registrate: {len(mods)}")
    print(f"   • Glossario aggiornato: {len(glossary)} termini")
    print(f"   • Success rate: {total_corrections/total_paragraphs:.1%}")
    
    # Mostra il glossario
    if glossary:
        print(f"   • Glossario: {', '.join(sorted(glossary))}")
    
    # Statistiche del safe corrector
    stats = safe_corrector.get_correction_stats()
    print(f"   • Safe corrector stats: {stats}")
    
    print("✅ Pipeline test completed successfully!")
    
    return doc, mods, glossary

async def test_error_handling():
    """Test della gestione errori"""
    print("\n=== TEST GESTIONE ERRORI ===")
    
    safe_corrector = SafeCorrector(conservative_mode=True, quality_threshold=0.85)
    
    # Mock paragraph
    class MockParagraph:
        def __init__(self, text):
            self.text = text
    
    # Test con funzione che genera errore
    def failing_correction(text):
        raise Exception("Simulated correction failure")
    
    para = MockParagraph("Test text")
    result = safe_corrector.correct_with_rollback(para, failing_correction, "error_test")
    
    print(f"❌ Error handling: applied={result.applied}, reason='{result.rollback_reason}'")
    assert not result.applied
    assert result.rollback_reason and "Function execution failed" in result.rollback_reason
    
    # Test con correzione che ritorna contenuto pericoloso
    def dangerous_correction(text):
        return "Completely different text that destroys content"
    
    result2 = safe_corrector.correct_with_rollback(para, dangerous_correction, "dangerous_test")
    print(f"🛡️  Dangerous correction blocked: applied={result2.applied}, quality={result2.quality_score.overall_score:.1%}")
    assert not result2.applied
    
    print("✅ Error handling test completed!")

if __name__ == "__main__":
    print("🧪 TESTING SAFE CORRECTION PIPELINE")
    print("=" * 60)
    
    async def run_all_tests():
        try:
            await test_safe_pipeline()
            await test_error_handling()
            
            print("\n" + "=" * 60)
            print("🎉 TUTTI I TEST DELLA PIPELINE COMPLETATI CON SUCCESSO!")
            print("✅ Sistema di correzione sicura integrato e funzionante")
            
        except Exception as e:
            print(f"\n❌ TEST FALLITO: {e}")
            raise
    
    asyncio.run(run_all_tests())
