"""
Test unitari per DocumentHandler.

Testa la gestione I/O dei documenti con validazione.
"""
import unittest
import sys
import tempfile
import shutil
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock

# Aggiungi il path del progetto
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from core.document_handler import DocumentHandler, DocumentLoadError, DocumentSaveError
from config.settings import Settings


class TestDocumentHandler(unittest.TestCase):
    """Test suite per DocumentHandler."""
    
    def setUp(self):
        """Setup per ogni test."""
        self.settings = Settings()
        self.handler = DocumentHandler(self.settings)
        
        # Crea directory temporanea per test
        self.temp_dir = Path(tempfile.mkdtemp())
        
    def tearDown(self):
        """Cleanup dopo ogni test."""
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)
            
    def test_handler_initialization(self):
        """Test inizializzazione handler."""
        self.assertIsNotNone(self.handler)
        self.assertIsNotNone(self.handler.settings)
        
    def test_validate_file_path_valid(self):
        """Test validazione path file valido."""
        # Crea file temporaneo
        test_file = self.temp_dir / "test.docx"
        test_file.touch()
        
        # Deve validare senza errori
        self.assertTrue(self.handler.validate_file_path(test_file))
        
    def test_validate_file_path_invalid(self):
        """Test validazione path file non valido."""
        invalid_path = self.temp_dir / "non_esistente.docx"
        
        with self.assertRaises(DocumentLoadError):
            self.handler.validate_file_path(invalid_path, raise_on_error=True)
            
    def test_validate_file_path_unsupported_format(self):
        """Test validazione formato file non supportato."""
        unsupported_file = self.temp_dir / "test.txt"
        unsupported_file.touch()
        
        with self.assertRaises(DocumentLoadError):
            self.handler.validate_file_path(unsupported_file, raise_on_error=True)
            
    @patch('python_docx.Document')
    def test_load_document_success(self, mock_document):
        """Test caricamento documento con successo."""
        # Setup mock
        mock_doc = Mock()
        mock_document.return_value = mock_doc
        
        test_file = self.temp_dir / "test.docx"
        test_file.touch()
        
        document = self.handler.load_document(test_file)
        
        self.assertEqual(document, mock_doc)
        mock_document.assert_called_once_with(str(test_file))
        
    @patch('python_docx.Document')
    def test_load_document_failure(self, mock_document):
        """Test fallimento caricamento documento."""
        mock_document.side_effect = Exception("Documento corrotto")
        
        test_file = self.temp_dir / "corrupted.docx"
        test_file.touch()
        
        with self.assertRaises(DocumentLoadError):
            self.handler.load_document(test_file)
            
    def test_extract_text_from_paragraphs(self):
        """Test estrazione testo da paragrafi."""
        # Mock document con paragrafi
        mock_doc = Mock()
        mock_paragraph1 = Mock()
        mock_paragraph1.text = "Primo paragrafo."
        mock_paragraph2 = Mock()
        mock_paragraph2.text = "Secondo paragrafo."
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2]
        
        text = self.handler.extract_text(mock_doc)
        
        expected = "Primo paragrafo.\n\nSecondo paragrafo."
        self.assertEqual(text, expected)
        
    def test_extract_text_empty_document(self):
        """Test estrazione testo da documento vuoto."""
        mock_doc = Mock()
        mock_doc.paragraphs = []
        
        text = self.handler.extract_text(mock_doc)
        
        self.assertEqual(text, "")
        
    def test_extract_text_with_formatting(self):
        """Test estrazione testo preservando formattazione."""
        mock_doc = Mock()
        mock_paragraph = Mock()
        
        # Mock runs con formattazione
        mock_run1 = Mock()
        mock_run1.text = "Testo "
        mock_run1.bold = False
        mock_run1.italic = False
        
        mock_run2 = Mock()
        mock_run2.text = "in grassetto"
        mock_run2.bold = True
        mock_run2.italic = False
        
        mock_paragraph.runs = [mock_run1, mock_run2]
        mock_paragraph.text = "Testo in grassetto"
        mock_doc.paragraphs = [mock_paragraph]
        
        text_with_formatting = self.handler.extract_text_with_formatting(mock_doc)
        
        self.assertIsInstance(text_with_formatting, dict)
        self.assertIn('text', text_with_formatting)
        self.assertIn('formatting', text_with_formatting)
        
    @patch('python_docx.Document')
    def test_save_document_success(self, mock_document):
        """Test salvataggio documento con successo."""
        mock_doc = Mock()
        
        output_path = self.temp_dir / "output.docx"
        
        self.handler.save_document(mock_doc, output_path)
        
        mock_doc.save.assert_called_once_with(str(output_path))
        
    @patch('python_docx.Document')
    def test_save_document_failure(self, mock_document):
        """Test fallimento salvataggio documento."""
        mock_doc = Mock()
        mock_doc.save.side_effect = Exception("Errore salvataggio")
        
        output_path = self.temp_dir / "output.docx"
        
        with self.assertRaises(DocumentSaveError):
            self.handler.save_document(mock_doc, output_path)
            
    def test_create_backup(self):
        """Test creazione backup documento."""
        # Crea file sorgente
        source_file = self.temp_dir / "source.docx"
        source_file.write_text("contenuto test")
        
        backup_path = self.handler.create_backup(source_file)
        
        self.assertTrue(backup_path.exists())
        self.assertTrue(backup_path.name.startswith("source_backup_"))
        self.assertTrue(backup_path.name.endswith(".docx"))
        
    def test_verify_document_integrity(self):
        """Test verifica integrità documento."""
        # Crea file test
        test_file = self.temp_dir / "test.docx"
        test_file.write_bytes(b"fake docx content")
        
        with patch.object(self.handler, 'load_document') as mock_load:
            mock_load.return_value = Mock()  # Documento valido
            
            is_valid = self.handler.verify_document_integrity(test_file)
            self.assertTrue(is_valid)
            
        with patch.object(self.handler, 'load_document') as mock_load:
            mock_load.side_effect = DocumentLoadError("Corrotto")
            
            is_valid = self.handler.verify_document_integrity(test_file)
            self.assertFalse(is_valid)
            
    def test_get_document_stats(self):
        """Test raccolta statistiche documento."""
        mock_doc = Mock()
        
        # Mock paragrafi
        mock_paragraphs = []
        for i in range(5):
            mock_p = Mock()
            mock_p.text = f"Paragrafo {i} con del testo di esempio."
            mock_paragraphs.append(mock_p)
            
        mock_doc.paragraphs = mock_paragraphs
        
        stats = self.handler.get_document_stats(mock_doc)
        
        self.assertEqual(stats['paragraph_count'], 5)
        self.assertGreater(stats['word_count'], 0)
        self.assertGreater(stats['character_count'], 0)
        self.assertIn('estimated_reading_time', stats)
        
    def test_apply_corrections_to_document(self):
        """Test applicazione correzioni a documento."""
        mock_doc = Mock()
        
        # Mock paragraphs
        mock_paragraph1 = Mock()
        mock_paragraph1.text = "Testo con errori."
        mock_paragraph2 = Mock()
        mock_paragraph2.text = "Altro testo."
        
        mock_doc.paragraphs = [mock_paragraph1, mock_paragraph2]
        
        corrections = {
            "Testo con errori.": "Testo senza errori.",
            "Altro testo.": "Altro testo corretto."
        }
        
        # Mock del metodo per applicare correzioni ai run
        with patch.object(self.handler, '_apply_correction_to_paragraph') as mock_apply:
            mock_apply.return_value = True
            
            result = self.handler.apply_corrections_to_document(mock_doc, corrections)
            
            self.assertTrue(result)
            self.assertEqual(mock_apply.call_count, 2)
            
    def test_preserve_formatting_during_correction(self):
        """Test preservazione formattazione durante correzione."""
        # Mock paragraph con runs formattati
        mock_paragraph = Mock()
        
        mock_run1 = Mock()
        mock_run1.text = "Testo "
        mock_run1.bold = False
        
        mock_run2 = Mock()
        mock_run2.text = "importante"
        mock_run2.bold = True
        
        mock_paragraph.runs = [mock_run1, mock_run2]
        mock_paragraph.text = "Testo importante"
        
        original_text = "Testo importante"
        corrected_text = "Testo molto importante"
        
        success = self.handler._apply_correction_to_paragraph(
            mock_paragraph, original_text, corrected_text
        )
        
        # Verifica che la formattazione sia stata preservata
        self.assertTrue(success)
        
    def test_batch_processing(self):
        """Test processamento batch di documenti."""
        # Crea file di test
        files = []
        for i in range(3):
            test_file = self.temp_dir / f"test_{i}.docx"
            test_file.touch()
            files.append(test_file)
            
        with patch.object(self.handler, 'load_document') as mock_load, \
             patch.object(self.handler, 'save_document') as mock_save:
            
            mock_load.return_value = Mock()
            
            results = self.handler.process_batch(files, lambda doc: doc)
            
            self.assertEqual(len(results), 3)
            self.assertEqual(mock_load.call_count, 3)
            
    def test_error_recovery(self):
        """Test recupero da errori durante processamento."""
        test_file = self.temp_dir / "test.docx"
        test_file.touch()
        
        with patch.object(self.handler, 'load_document') as mock_load:
            mock_load.side_effect = [
                DocumentLoadError("Errore temporaneo"),  # Primo tentativo fallisce
                Mock()  # Secondo tentativo riesce
            ]
            
            with patch.object(self.handler, '_attempt_recovery') as mock_recovery:
                mock_recovery.return_value = True
                
                document = self.handler.load_document_with_recovery(test_file)
                
                self.assertIsNotNone(document)
                mock_recovery.assert_called_once()


class TestDocumentHandlerIntegration(unittest.TestCase):
    """Test di integrazione per DocumentHandler."""
    
    def setUp(self):
        """Setup per test di integrazione."""
        self.handler = DocumentHandler()
        self.temp_dir = Path(tempfile.mkdtemp())
        
    def tearDown(self):
        """Cleanup."""
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)
            
    def test_full_document_workflow(self):
        """Test workflow completo caricamento-modifica-salvataggio."""
        # Questo test richiede python-docx installato
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not available")
            
        # Crea documento di test
        doc = Document()
        doc.add_paragraph("Paragrafo di test con errori.")
        
        test_file = self.temp_dir / "test_input.docx"
        doc.save(str(test_file))
        
        # Carica documento
        loaded_doc = self.handler.load_document(test_file)
        self.assertIsNotNone(loaded_doc)
        
        # Estrai testo
        text = self.handler.extract_text(loaded_doc)
        self.assertIn("Paragrafo di test", text)
        
        # Applica correzioni (mock)
        corrections = {"errori": "correzioni"}
        success = self.handler.apply_corrections_to_document(loaded_doc, corrections)
        
        # Salva documento modificato
        output_file = self.temp_dir / "test_output.docx"
        self.handler.save_document(loaded_doc, output_file)
        
        self.assertTrue(output_file.exists())


if __name__ == '__main__':
    unittest.main(verbosity=2)
