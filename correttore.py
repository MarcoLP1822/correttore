#correttore.py
"""
Corregge grammatica e ortografia in **qualsiasi** parte di un documento word (.docx) e produce un report Markdown con tutte le modifiche.

* testo normale, anche in tabelle nidificate
* header & footer
* testo delle note a piè di pagina (footnotes)
* caselle di testo / forme (<w:txbxContent>)

Mantiene corsivi, grassetti, sottolineature, riferimenti di nota, ecc.
Richiede **python-docx ≥ 0.8.11**.
"""

from __future__ import annotations

# ——— Standard library ————————————————————————————————
import logging

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger(__name__)
logging.getLogger("urllib3").setLevel(logging.WARNING)

import asyncio
import aiofiles
import collections
import json
import os
import re
import shutil
import time
import zipfile
from pathlib import Path
from copy import deepcopy
from datetime import datetime
from precheck import has_errors
from dataclasses import dataclass
from difflib import SequenceMatcher
from typing import Dict, Iterable, List, Optional, Tuple
from pprint import pprint

# ——— Third-party ————————————————————————————————
import tiktoken
from tqdm import tqdm
from tqdm.asyncio import tqdm_asyncio
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree
from lxml.etree import XMLSyntaxError
from openai import AsyncOpenAI, OpenAI
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor


# Carica le variabili d'ambiente dal file .env.local
load_dotenv('.env.local')

# ——— Local modules ————————————————————————————————
from reports import write_glossary_report, write_markdown_report
from token_utils import tokenize, token_starts, count_tokens
from settings import OPENAI_MODEL, MAX_TOKENS
from openai_client import get_async_client
from utils_openai import get_corrections_async, build_messages
from spellfix import spellfix_paragraph
from grammarcheck import grammarcheck
from llm_correct import llm_correct, llm_correct_batch

LT_POOL = ProcessPoolExecutor(max_workers=os.cpu_count() - 1 or 1)

# ───────────────────────── CONFIGURAZIONE ────────────────────────────
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

try:
    ENC = tiktoken.encoding_for_model(OPENAI_MODEL)
except KeyError:
    # fallback universale, compatibile con GPT-4/3.5
    ENC = tiktoken.get_encoding("cl100k_base")
# ─────────────────────────────────────────────────────────────────────

from token_utils import NAME_RE, WORD_RE

GLOSSARY_STOP = {
    "CAPITOLO", "CAPITOLI", "PROLOGO", "EPILOGO",
    "INDICE", "RINGRAZIAMENTI", "DEDICA", "DEDICHE", "PREFAZIONE", "PREFAZIONI",
}
# ─────────────────── Pandoc: ODT/DOC → DOCX ──────────────────────────
import win32com.client as win32

def convert_to_docx(src: Path) -> Path:
    if src.suffix.lower() == ".docx":
        return src

    dst = src.with_suffix(".docx")
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(str(src))
        doc.SaveAs(str(dst), FileFormat=16)   # 16 = wdFormatDocumentDefault
        doc.Close()
    finally:
        word.Quit()
    return dst

# ──────────────────────────────────────────────────────────────────────

# ╭─────────────────────────── Note a piè di pagina ▾──────────────────────────╮
from collections import defaultdict

async def correggi_footnotes_xml_async(docx_path: Path,
                                       async_client: AsyncOpenAI,
                                       glossary: set[str] | None = None) -> None:

    """
    Corregge le note a piè di pagina contenute in word/footnotes.xml
    (refusi, ortografia, punteggiatura) preservando la formattazione
    run-per-run del documento Word.
    """
    glossary = glossary or set()           # se None, usa set vuoto
    tmp_dir = "_tmp_docx_unzipped"

    # 1) Estrai il .docx in una cartella temporanea --------------------
    with zipfile.ZipFile(docx_path, "r") as zf:
        zf.extractall(tmp_dir)

    footnotes_file = os.path.join(tmp_dir, "word", "footnotes.xml")
    if not os.path.exists(footnotes_file):
        shutil.rmtree(tmp_dir)
        return                              # documento senza note
    
    # Se il file esiste ma è vuoto, non ci sono note → esci
    if os.path.getsize(footnotes_file) == 0:
        shutil.rmtree(tmp_dir)
        return

    # Prova a fare il parse; se fallisce, esci comunque in sicurezza
    try:
        ns   = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        tree = etree.parse(footnotes_file)
    except XMLSyntaxError:
        shutil.rmtree(tmp_dir)
        return

    ns   = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    tree = etree.parse(footnotes_file)
    footnotes = tree.xpath(
        "//w:footnote[not(@w:type='separator')]",
        namespaces=ns
    )

# 3) Ciclo con barra di avanzamento ------------------------------------
    for foot in tqdm(
            footnotes,
            desc="Note a piè di pagina",
            total=len(footnotes)
    ):
        txt_nodes = foot.xpath(".//w:t", namespaces=ns)
        full_text = "".join(n.text or "" for n in txt_nodes)

        if not full_text.strip():
            continue                             # nota vuota → salta

    # 3a) Chiamata OpenAI (async) --------------------------------------
        raw = await get_corrections_async(
            payload_json=json.dumps(
                [{"id": 0, "txt": full_text}],
                ensure_ascii=False
            ),
            client=async_client,
            glossary=glossary,
            context="",
        )
        if isinstance(raw, list):
            corr_list = raw 
        elif isinstance(raw, dict) and "corr" in raw:
            corr_list = raw["corr"]
        else:
            raise TypeError(
                f"Formato inatteso da get_corrections_async(): {type(raw)} – {raw!r}"
        )

        corrected = corr_list[0]["txt"]

        if corrected == full_text:
            continue                             # nessuna correzione

    # 3b) Redistribuzione token nei nodi <w:t> ... (resto del tuo codice)
    # ------------------------------------------------------------------


        # 2b) Redistribuisci token corretti nei singoli <w:t> ----------
        orig_tok   = tokenize(full_text)
        corr_tok   = tokenize(corrected)
        mapping    = align_tokens(orig_tok, corr_tok)
        starts_orig = token_starts(orig_tok)

        # mappa carattere→indice nodo
        char2node = []
        for idx, n in enumerate(txt_nodes):
            char2node.extend([idx] * len(n.text or ""))

        tok_per_node = defaultdict(list)
        for ref_idx, tok in mapping:
            if ref_idx is None or ref_idx >= len(starts_orig):
                node_idx = 0
            else:
                char_pos = starts_orig[ref_idx]
                node_idx = char2node[min(char_pos, len(char2node) - 1)]
            tok_per_node[node_idx].append(tok)

        # scrivi il testo nei nodi preservando la partizione originale
        for idx, n in enumerate(txt_nodes):
            n.text = "".join(tok_per_node.get(idx, []))

        # 2c) Se si perde la formattazione del primo run, ristabiliscila
        first_with_txt = next((n for n in txt_nodes if n.text and n.text.strip()), None)
        if first_with_txt is not None:
            run_first       = first_with_txt.getparent()
            has_rPr_first   = run_first.find("./w:rPr", namespaces=ns)
            has_bold_first  = run_first.xpath("./w:rPr/w:b", namespaces=ns)

            if not has_bold_first:
                for n2 in txt_nodes:
                    if n2 is first_with_txt or not (n2.text and n2.text.strip()):
                        continue
                    rPr_other = n2.getparent().find("./w:rPr", namespaces=ns)
                    if rPr_other is not None and list(rPr_other):
                        if has_rPr_first is None:
                            has_rPr_first = etree.SubElement(run_first, qn("w:rPr"))
                        for child in rPr_other:
                            has_rPr_first.append(deepcopy(child))
                        break   # una sola copia è sufficiente

    # 3) Salva il nuovo footnotes.xml e ricompatta il .docx -------------
    async with aiofiles.open(footnotes_file, "wb") as f:
        await f.write(etree.tostring(tree,
                           xml_declaration=True,
                           encoding="utf-8"))

    tmp_docx = docx_path.with_suffix(".tmp")
    with zipfile.ZipFile(tmp_docx, "w") as zf:
        for root, _, files in os.walk(tmp_dir):
            for f in files:
                fullpath = os.path.join(root, f)
                arcname  = os.path.relpath(fullpath, tmp_dir)
                zf.write(fullpath, arcname)

    shutil.move(tmp_docx, docx_path)        # sovrascrive l'originale
    shutil.rmtree(tmp_dir)
    logger.info("✏️  Note a piè di pagina corrette")
# ╭─────────────────────────── Data-model modifiche ▾────────────────────╮
@dataclass
class Modification:
    par_id: int
    original: str
    corrected: str
# ╰──────────────────────────────────────────────────────────────────────╯
# ───────────────────── NEW universal chunker ──────────────────────
def chunk_paragraphs_all(
    paragraphs: List[Paragraph],
    max_tokens: int = 3_000,
    max_pars:   int = 5,
) -> List[List[Paragraph]]:
    """
    Divide *tutti* i paragrafi in blocchi:
      • al massimo `max_pars` paragrafi
      • e al massimo `max_tokens` token complessivi
    Ritorna una lista di chunk (liste di Paragraph).
    """
    chunks: List[List[Paragraph]] = []
    current, current_tok = [], 0
    for p in paragraphs:
        tok = count_tokens(p.text)
        # se il singolo paragrafo supera il limite token, lo isolo
        if tok > max_tokens:
            if current:
                chunks.append(current)
                current, current_tok = [], 0
            chunks.append([p])
            continue
        # se aggiungerlo sfora i limiti → chiudi chunk e apri nuovo
        if (current and
            (current_tok + tok > max_tokens or len(current) >= max_pars)):
            chunks.append(current)
            current, current_tok = [], 0
        current.append(p)
        current_tok += tok
    if current:
        chunks.append(current)
    return chunks
# ───────────────────────────────────────────────────────────────────

# ╭─────────────────────────── Chunking ▾────────────────────────────────╮
#def chunk_paragraph_objects(
#    paragraphs: List[Paragraph],
#    max_tokens: int = MAX_TOKENS,
#) -> List[List[Paragraph]]:
#    """Dividi la lista di oggetti Paragraph in blocchi < max_tokens."""
#    chunks: List[List[Paragraph]] = []
#    current: List[Paragraph] = []
#    current_tokens = 0

#    for p in paragraphs:
#        para_tokens = count_tokens(p.text)

        # 1) se il singolo paragrafo sfora il limite, isolalo
#        if para_tokens > max_tokens:
#            if current:                     # salva il chunk in corso
#                chunks.append(current)
#                current, current_tokens = [], 0
#            chunks.append([p])              # paragrafo “oversize” da solo
#            continue                        # salta al paragrafo successivo

#        # 2) comportamento normale
#        if current and current_tokens + para_tokens > max_tokens:
#            chunks.append(current)
#            current = [p]
#            current_tokens = para
#        else:
#            current.append(p)
#            current_tokens += para_tokens

#    if current:
#        chunks.append(current)
#    return chunks
# ╰──────────────────────────────────────────────────────────────────────╯

# ╭─────────────────────────── Diff token-level ▾────────────────────────╮
def align_tokens(orig: List[str], corr: List[str]) -> List[Tuple[Optional[int], str]]:
    sm = SequenceMatcher(a=orig, b=corr, autojunk=False)
    out: List[Tuple[Optional[int], str]] = []
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == "equal":
            for k in range(i2 - i1):
                out.append((i1 + k, corr[j1 + k]))
        elif op == "replace":
            inherit = i1 if i1 < len(orig) else None
            for k in range(j2 - j1):
                out.append((inherit, corr[j1 + k]))
        elif op == "insert":
            inherit = i1 - 1 if i1 > 0 else None
            for k in range(j1, j2):
                out.append((inherit, corr[k]))
    return out
# ╰──────────────────────────────────────────────────────────────────────╯

# ╭─────────────────────────── Helpers copy RUN ▾────────────────────────╮
def copy_rPr(src_rPr, dest_run):
    if src_rPr is None:
        return
    dest = dest_run._r.get_or_add_rPr()
    dest.clear()
    for child in src_rPr:
        dest.append(deepcopy(child))


def clone_run(src_run, paragraph):
    new_run = paragraph.add_run("")
    copy_rPr(src_run._r.rPr, new_run)
    for child in src_run._r:
        if child.tag != qn("w:t"):
            new_run._r.append(deepcopy(child))
    return new_run
# ╰──────────────────────────────────────────────────────────────────────╯

# ───────────────── helper “quanto è diverso?” ────────────────────────
# ─────────────────────────────────────────────────────────────────────
def _too_much_cut(orig: str, corr: str,
                  lost_thresh: float = 0.15) -> bool:
    """
    True **solo** se l’output:
      • perde >10 % dei token **e**
      • ha meno frasi dell’originale
    (Quasi impossibile in una correzione normale)
    """
    tok_orig = tokenize(orig)
    tok_corr = tokenize(corr)
    lost = (len(tok_orig) - len(tok_corr)) / max(1, len(tok_orig))
    if lost <= lost_thresh:
        return False                         # taglio contenuto: ok

    # conteggio frasi
    sent = lambda s: len([t for t in re.split(r"[.!?…]", s) if t.strip()])
    return sent(corr) < sent(orig)

# ╭──────────────────────── Nuova versione step-safe ───────────────────╮
def _run_map(paragraph: Paragraph) -> list[int]:
    """Mappa ogni carattere al run di provenienza (es. [0,0,0,1,1,2,…])."""
    m = []
    for idx, r in enumerate(paragraph.runs):
        m.extend([idx] * len(r.text))
    return m

def apply_correction_to_paragraph(
    p: Paragraph,
    corrected: str,
    mods: list[Modification],
    par_id: int,
    glossary: set[str],
):
    original = p.text
    # 1) quote  +  line-break(s)  →  quote + space
    corrected = re.sub(r'([«“"”])\s*[\r\n]+\s*', r'\1 ', corrected)

    # 2) line-break(s)  +  quote  →  space + quote
    corrected = re.sub(r'[\r\n]+\s*([«“"”])', r' \1', corrected)
    
    # 3) se nel testo originale non c'erano line-break interni, eliminali tutti
    if ("\n" in corrected or "\r" in corrected) and ("\n" not in original and "\r" not in original):
        corrected = corrected.replace("\r", " ").replace("\n", " ")

    if corrected == original:            
        return

    # ── filtro sicurezza minimo (problema B, v. § 2 per il resto) ─────
    if _too_much_cut(original, corrected):
        return

    # diff token-level → quali token cambiano?
    orig_tok = tokenize(original)
    corr_tok = tokenize(corrected)
    mapping  = align_tokens(orig_tok, corr_tok)
    if all(a == b for a, b in zip(orig_tok, corr_tok)):
        return                                # cambi trascurabili

    mods.append(Modification(par_id, original, corrected))

    # 1) prepara struttura dati
    char2run = _run_map(p)
    tok_per_run: dict[int, list[str]] = defaultdict(list)
    starts_orig = token_starts(orig_tok)

    # 2) assegna i token corretti al run “ereditato”
    last_idx = 0
    for ref_idx, tok in mapping:
        if ref_idx is None:
            run_idx = last_idx
        else:
            pos = starts_orig[ref_idx]
            run_idx = char2run[min(pos, len(char2run)-1)]
        tok_per_run[run_idx].append(tok)
        last_idx = run_idx

    # 3) sostituisci solo i run che contengono testo (salvi stili & ponct.)
    for idx, run in enumerate(p.runs):
        if idx in tok_per_run:
            run.text = "".join(tok_per_run[idx])
    
    # 4) post-fix: accorpa run che contengono solo virgolette -------------
    orphan_quotes = {'"', '«', '“', '”'}
    runs = list(p.runs)
    for i, run in enumerate(runs):
        t = (run.text or "").strip()
        if t in orphan_quotes:
            # se c'è un run successivo con testo, fondi la virgoletta all'inizio
            if i + 1 < len(runs) and (runs[i + 1].text or "").strip():
                runs[i + 1].text = run.text + runs[i + 1].text
                run.text = ""           # svuota il run-virgoletta

    # 5) rimuovi il run che contiene solo <w:br/> dopo una virgoletta ------
    runs = list(p.runs)                      # lista aggiornata
    open_quotes = {'"', '«', '“', '”', '‘', '’'}
    for i in range(len(runs) - 1):
        prev_txt = (runs[i].text or "")
        this_txt = (runs[i + 1].text or "")
        has_only_br = not this_txt.strip() and runs[i + 1]._r.xpath("./w:br")
        if prev_txt and prev_txt[-1] in open_quotes and has_only_br:
            # elimina il run con solo <w:br/>
            parent = runs[i + 1]._r.getparent()
            parent.remove(runs[i + 1]._r)

    # 6) aggiorna dinamicamente il glossario
    for name in NAME_RE.findall(corrected):
        if name.upper() not in GLOSSARY_STOP:
            glossary.add(name)
# ╰─────────────────────────────────────────────────────────────────────╯

async def grammarcheck_async(text: str) -> str:
    """
    Esegue grammarcheck() in un thread del pool,
    così possiamo processare paragrafi in parallelo.
    """
    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(LT_POOL, grammarcheck, text)

def _minor_change(a: str, b: str) -> bool:
    import unicodedata, re
    norm = lambda s: re.sub(r"\W+", "", unicodedata.normalize("NFD", s).lower())
    return norm(a) == norm(b)

# ─────────── Pipeline locale + batch GPT multiparagrafo ───────────

# ─────────────────────────── correct_paragraph_group ──────────────────────────
async def correct_paragraph_group(
    paragraphs:   list[Paragraph],
    all_paras:    list[Paragraph],   # non usato qui, ma lasciato per compatibilità
    start_par_id: int,
    client:       AsyncOpenAI,
    glossary:     set[str],
    mods:         list[Modification],
    context_size: int = 3,
):
    """
    • Spell-check + grammar-check su ogni paragrafo.
    • Se il testo NON cambia → lo accoda per la correzione GPT in batch.
    • Una sola chiamata GPT per chunk; cache usata automaticamente.
    • Log dettagliati per capire tempi e cache.
    """
    logger = logging.getLogger("PAR-GROUP")
    logger.debug("Chunk START  paragrafi=%d  start_id=%d", len(paragraphs), start_par_id)

    # liste per i paragrafi che servono GPT
    pending_idx: list[int] = []
    pending_txt: list[str] = []

    # ── Passo 1: correzione locale (spellfix + grammarcheck) ─────────────
    for idx, p in enumerate(paragraphs):
        original = p.text
        if not original.strip():
            continue

        step1 = spellfix_paragraph(original, glossary)
        step2 = await grammarcheck_async(step1)

        if step2 != original:
            if _minor_change(original, step2):
                # differenza microscopica → lascia che GPT lo riveda
                pending_idx.append(idx)
                pending_txt.append(step2)
            else:
                apply_correction_to_paragraph(
                    p, step2, mods, start_par_id + idx, glossary
                )
        else:
            # nessuna modifica locale → manda al batch GPT
            pending_idx.append(idx)
            pending_txt.append(step2)

    logger.info(
        "Chunk %d: local=%d  GPT=%d",
        start_par_id,
        len(paragraphs) - len(pending_idx),
        len(pending_idx),
    )

    # ── Passo 2: se nulla richiede GPT, finito ───────────────────────────
    if not pending_txt:
        logger.debug("Chunk %d finito (nessuna chiamata GPT).", start_par_id)
        return

    # ── Passo 3: correzione batch con GPT (usa cache interna) ────────────
    corrected_list = await llm_correct_batch(pending_txt, client)

    # ── Passo 4: applica le correzioni ricevute dal batch GPT ────────────
    for local_i, new_text in zip(pending_idx, corrected_list):
        p = paragraphs[local_i]
        if new_text != p.text:
            apply_correction_to_paragraph(
                p, new_text, mods,
                start_par_id + local_i,
                glossary,
            )
            logger.debug("✓ GPT fix     id=%d", start_par_id + local_i)

    logger.debug("Chunk END %d", start_par_id)
# ───────────────────────────────────────────────────────────────────────────────

# ╭─────────────────── Funzione generica di correzione ─────────────────╮
#async def correct_paragraph_group(
#    paragraphs:   list[Paragraph],
#    all_paras:    list[Paragraph],
#    start_par_id: int,
#    client:       AsyncOpenAI,
#    glossary:     set[str],
#    mods:         list[Modification],
#    context_size: int = 3,
#):
#    sem = asyncio.Semaphore(10)  # max 10 chiamate in parallelo

#    async def process_one(par_id: int, p: Paragraph) -> tuple[Paragraph, str]:
#        original_text = p.text
#        if not original_text.strip():
#            return p, original_text

#        step1 = spellfix_paragraph(original_text, glossary)
#        step2 = await grammarcheck_async(step1)

        # Solo se spellfix+grammarcheck NON hanno cambiato nulla → LLM
#            if step2 == original_text:
#                async with sem:
#                    final = await llm_correct(step2, client)
#            else:
#                final = step2


#            return p, final

#        tasks = [
#            asyncio.create_task(process_one(start_par_id + i, p))
#            for i, p in enumerate(paragraphs)
#        ]

#        results = await asyncio.gather(*tasks)

#        for i, (p, corrected_text) in enumerate(results):
#            if corrected_text != p.text:
#                apply_correction_to_paragraph(
#                    p,
#                    corrected_text,
#                    mods,
#                    start_par_id + i,
#                    glossary,
#                )
# ╰──────────────────────────────────────────────────────────────────────╯
# ╭──────────────────── Wrapper: corpo principale ─────────────────────╮
async def fix_body_chunks(
    async_client: AsyncOpenAI,
    all_paras:    list[Paragraph],
    para_chunks:  list[list[Paragraph]],
    start_id:     int,
    mods:         list[Modification],
    glossary:     set[str],
):
    """
    Scorre tutti i chunk creati da `chunk_paragraph_objects`, lancia
    in parallelo `correct_paragraph_group` e mostra una progress-bar
    che avanza quando ciascun chunk è realmente completato.
    """
    tasks: list[asyncio.Task] = []
    par_id = start_id

    # Costruisci la lista dei task mantenendo l'ID di paragrafo corretto
    for chunk in para_chunks:
        tasks.append(
            correct_paragraph_group(
                paragraphs   = chunk,
                all_paras    = all_paras,
                start_par_id = par_id,
                client       = async_client,
                glossary     = glossary,
                mods         = mods,
            )
        )
        par_id += len(chunk)

    # Progress-bar: si aggiorna al termine di ogni task
    await tqdm_asyncio.gather(
        *tasks,
        total=len(tasks),
        desc="Correzione chunk"
    )
# ╰──────────────────────────────────────────────────────────────────────╯

# ╭─────────────────────────── Map char→run ▾────────────────────────────╮
def char_to_run_map(paragraph) -> List[int]:
    mapping: List[int] = []
    for idx, run in enumerate(paragraph.runs):
        mapping.extend([idx] * len(run.text))
    return mapping
# ╰──────────────────────────────────────────────────────────────────────╯
# ───────────────────── Detect equazioni OMML ──────────────────────────
def has_math(paragraph: Paragraph) -> bool:
    """
    Ritorna True se il <w:p> contiene almeno un elemento <m:oMath>
    o <m:oMathPara> (= equazioni Word).
    """
    return bool(
        paragraph._p.xpath(
            ".//*[local-name()='oMath' or local-name()='oMathPara']"
        )
    )
# ──────────────────────────────────────────────────────────────────────
# ───────────────────────── traversal del documento ─────────────────────────────
def iter_body_paragraphs(container) -> Iterable[Paragraph]:
    for para in container.paragraphs:
        yield para
    for tbl in getattr(container, "tables", []):
        for row in tbl.rows:
            for cell in row.cells:
                yield from iter_body_paragraphs(cell)


def iter_footnote_paragraphs(doc: Document) -> Iterable[Paragraph]:
    fpart = getattr(doc.part, "footnotes_part", None)
    if fpart:
        for footnote in fpart.footnotes:
            for para in footnote.paragraphs:
                yield para


def iter_header_footer_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for sect in doc.sections:
        for hf in (sect.header, sect.footer):
            if hf:
                yield from iter_body_paragraphs(hf)


def iter_textbox_paragraphs(doc: Document) -> Iterable[Paragraph]:
    parts = [doc.part]
    for sect in doc.sections:
        parts.extend([sect.header.part, sect.footer.part])
    fpart = getattr(doc.part, "footnotes_part", None)
    if fpart:
        parts.append(fpart)
    for part in parts:
        root = part._element
        for txbx in root.xpath('.//*[local-name()="txbxContent"]'):
            for p_el in txbx.xpath('.//*[local-name()="p"]'):
                yield Paragraph(p_el, part)


def iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    yield from iter_body_paragraphs(doc)
    yield from iter_footnote_paragraphs(doc)
    yield from iter_header_footer_paragraphs(doc)
    yield from iter_textbox_paragraphs(doc)
# ──────────────────────────────────────────────────────────────────────

# ───────────────── Helper: verifica se C è interamente contenuto in A ─────────
def _is_clone(text_a: str, text_c: str, threshold: float = 0.95) -> bool:
    """
    True se C è quasi uguale (≥95 %) alla coda di A.
    Usa SequenceMatcher per maggiore robustezza.
    """
    norm = lambda s: " ".join(s.split()).lower()
    return SequenceMatcher(None, norm(text_a), norm(text_c)).ratio() >= threshold
# ───────────────────────────────────────────────────────────────────────────────

# ───────────────── FINE DEGLI HELPER ────────────────
# ───────────────── Deduplica paragrafi clonati dopo un sectPr ────────────────
def remove_duplicate_after_sectpr(document: Document) -> None:
    """
    Se trova la sequenza:
        A = paragrafo con testo
        B = paragrafo VUOTO con <w:sectPr>
        C = paragrafo la cui *intera* stringa è già contenuta in coda ad A
    elimina C ma mantiene B (così l’interruzione di sezione resta).
    """
    paras = list(iter_body_paragraphs(document))
    for i in range(len(paras) - 2, 0, -1):
        A, B, C = paras[i - 1], paras[i], paras[i + 1]
        if B.text.strip():                       # B deve essere vuoto
            continue
        if not B._p.xpath("./w:pPr/w:sectPr"):   # B deve avere <w:sectPr>
            continue
        text_a = A.text.strip()
        text_c = C.text.strip()
        if text_c and _is_clone(text_a, text_c):
            C._p.getparent().remove(C._p)
# ─────────────────────────────────────────────────────────────────────────────

# ───────────────────────── entry-point con logging ────────────────────
def process_doc(inp: Path, out: Path):
    doc = Document(inp)

    # 1. Raccoglie tutti i paragrafi del documento (corpo, intestazioni, piè di pagina, note, caselle di testo)
    all_paras = list(iter_all_paragraphs(doc))
    paras_to_fix = [
        p
        for p in all_paras
        if has_errors(p.text) and not has_math(p)
    ]


    # 2. Crea un glossario iniziale con i nomi ricorrenti (apparsi almeno 2 volte)
    name_counts = collections.Counter(
        n for p in all_paras for n in NAME_RE.findall(p.text)
        if n.upper() not in GLOSSARY_STOP
    )
    glossary = {w for w, c in name_counts.items() if c >= 2}

    # 3. Suddivide il corpo del documento in chunk, rispettando il limite di token
    para_chunks = chunk_paragraphs_all(all_paras, max_tokens=3_000, max_pars=5)
    total_chunks = len(para_chunks)
    logger.info("🔍  Rilevati %s chunk universali (≤5 paragrafi, ≤3000 token).", total_chunks)

    # 4. Lista per raccogliere tutte le modifiche da riportare nel diff finale
    mods: list[Modification] = []

    # 5. Definisce la coroutine che gestisce l'intero flusso asincrono (corpo + note)
    async def handle_all(doc: Document, out_path: Path):
        async_client = get_async_client()

        # 5.1 Corregge il corpo del documento (in parallelo, chunk per chunk)
        await fix_body_chunks(
            async_client=async_client,
            all_paras=all_paras,
            para_chunks=para_chunks,
            start_id=1,
            mods=mods,
            glossary=glossary,
        )
        # Deduplica cloni post-correzione ma conserva l'interruzione di sezione
        remove_duplicate_after_sectpr(doc)  

        # 5.2 Salva il documento corretto prima di correggere le note (serve file .docx completo)
        doc.save(out_path)
        logger.info("💾  Documento salvato: %s", out_path.name)

        # 5.3 Corregge le note a piè di pagina in parallelo
        await correggi_footnotes_xml_async(
            docx_path=out_path,
            async_client=async_client,
            glossary=glossary,
        )

    # 6. Avvia l'evento asincrono completo
    asyncio.run(handle_all(doc, out))

    # 7. Genera i due report Markdown (diff + glossario)
    write_markdown_report(mods, out)
    write_glossary_report(glossary, all_paras, out)


# ───────────────────────────────────────────────────────────────────────
def find_latest_docx(folder: Path) -> Path:
    """
    Cerca il file più recente con estensione .docx, .odt o .doc (esclude i
    file di lock ~$.docx). Se non è già .docx lo converte con Pandoc.
    """
    candidates = [
        p for p in folder.iterdir()
        if p.is_file()
        and not p.name.startswith("~$")
        and p.suffix.lower() in (".docx", ".odt", ".doc")
    ]
    if not candidates:
        raise RuntimeError("Nessun documento .docx/.odt/.doc trovato nella cartella")

    latest = max(candidates, key=lambda p: p.stat().st_mtime)
    return convert_to_docx(latest)


if __name__ == "__main__":
    # marca inizio
    start_time = time.perf_counter()

    here = Path(__file__).resolve().parent
    src = find_latest_docx(here)
    dst = src.with_stem(src.stem + "_corretto")
    logger.info("📝  Correggo %s → %s …", src.name, dst.name)
    process_doc(src, dst)

    # tempo impiegato
    elapsed = time.perf_counter() - start_time
    logger.info("✨  Fatto in %.2f secondi!", elapsed)